import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings, HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.document_loaders import WebBaseLoader
import requests
from bs4 import BeautifulSoup
import urllib.parse
import pandas as pd
import openpyxl
import xlrd
from pydantic import BaseModel
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument
from langchain.prompts import PromptTemplate
from htmlTemplates import css, bot_template, user_template

import os
import warnings
import hashlib
import json
import shutil
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import logging
from pathlib import Path
import time

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Supprimer les avertissements
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Répertoires de stockage
VECTORSTORE_BASE_DIR = "vectorstores"
CACHE_DIR = "cache"

# Configuration des modèles
EMBEDDING_MODELS = {
    "OpenAI": {
        "class": OpenAIEmbeddings,
        "params": {},
        "requires_api_key": True
    },
    "HuggingFace Instructor": {
        "class": HuggingFaceInstructEmbeddings,
        "params": {"model_name": "hku-nlp/instructor-xl"},
        "requires_api_key": False
    },
    "HuggingFace Sentence Transformers": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/all-MiniLM-L6-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    },
    "HuggingFace BGE Small": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "BAAI/bge-small-en-v1.5",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    },
    "HuggingFace Multilingual": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    }
}

LLM_MODELS = {
    "Gemini 2.0 Flash": {
        "model_name": "gemini-2.0-flash-exp",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle puissant et rapide"
    },
    "Gemini Pro 1.5": {
        "model_name": "gemini-1.5-pro-latest",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini 1.5 Pro - Modèle le plus avancé"
    },
    "Gemini Pro": {
        "model_name": "gemini-pro",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle standard performant"
    }
}

class SegmentedVectorStoreManager:
    """Gestionnaire des vector stores FAISS avec segmentation par document"""
    
    def __init__(self, base_dir: str = VECTORSTORE_BASE_DIR):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(exist_ok=True)
        self.cache_dir = Path(CACHE_DIR)
        self.cache_dir.mkdir(exist_ok=True)
        
    def generate_vectorstore_id(self, documents: List[LangchainDocument], 
                              embedding_model: str, chunk_params: Dict) -> str:
        """Génère un ID unique pour le vector store basé sur le contenu et les paramètres"""
        content_hash = hashlib.md5()
        
        # Hash du contenu des documents
        for doc in documents:
            content_hash.update(doc.page_content.encode('utf-8'))
            # Inclure les métadonnées principales
            if 'source' in doc.metadata:
                content_hash.update(doc.metadata['source'].encode('utf-8'))
        
        # Hash des paramètres
        params_str = f"{embedding_model}_{chunk_params['chunk_size']}_{chunk_params['chunk_overlap']}"
        content_hash.update(params_str.encode('utf-8'))
        
        # Ajouter timestamp pour unicité
        timestamp = str(int(time.time()))
        content_hash.update(timestamp.encode('utf-8'))
        
        return content_hash.hexdigest()[:16]
    
    def get_vectorstore_path(self, vectorstore_id: str) -> Path:
        """Retourne le chemin du répertoire du vector store"""
        return self.base_dir / f"faiss_db_{vectorstore_id}"
    
    def vectorstore_exists(self, vectorstore_id: str) -> bool:
        """Vérifie si un vector store existe déjà"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        return vectorstore_path.exists() and (vectorstore_path / "index.faiss").exists()
    
    def save_vectorstore_metadata(self, vectorstore_id: str, metadata: Dict):
        """Sauvegarde les métadonnées du vector store"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        metadata_file = vectorstore_path / "metadata.json"
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, indent=2, ensure_ascii=False, fp=f)
    
    def load_vectorstore_metadata(self, vectorstore_id: str) -> Optional[Dict]:
        """Charge les métadonnées du vector store"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        metadata_file = vectorstore_path / "metadata.json"
        
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return None
    
    def get_document_segments(self, vectorstore_id: str) -> List[Dict]:
        """Récupère la liste des segments de documents disponibles"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        segments_dir = vectorstore_path / "segments"
        
        segments = []
        if segments_dir.exists():
            metadata = self.load_vectorstore_metadata(vectorstore_id)
            if metadata and metadata.get('segments'):
                for doc_id, segment_info in metadata['segments'].items():
                    segment_path = segments_dir / doc_id
                    if segment_path.exists() and (segment_path / "index.faiss").exists():
                        segments.append({
                            'doc_id': doc_id,
                            'source': segment_info.get('source', 'Unknown'),
                            'type': segment_info.get('type', 'unknown'),
                            'chunk_count': segment_info.get('chunk_count', 0),
                            'vector_count': segment_info.get('vector_count', 0),
                            'path': str(segment_path)
                        })
        
        return segments
    
    def save_segmented_vectorstore(self, vectorstore: FAISS, vectorstore_id: str, 
                                 documents: List[LangchainDocument], embedding_model: str,
                                 chunk_params: Dict, custom_name: str = None) -> bool:
        """Sauvegarde le vector store avec segmentation par document"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            # Supprimer le répertoire existant si nécessaire
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
            
            # Créer le répertoire
            vectorstore_path.mkdir(parents=True, exist_ok=True)
            
            # Sauvegarder le vector store principal
            vectorstore.save_local(str(vectorstore_path))
            
            # Créer des vector stores séparés pour chaque document source
            document_segments = self._create_document_segments(documents, embedding_model)
            
            # Sauvegarder chaque segment
            segments_dir = vectorstore_path / "segments"
            segments_dir.mkdir(exist_ok=True)
            
            for doc_id, segment_data in document_segments.items():
                segment_path = segments_dir / doc_id
                segment_path.mkdir(exist_ok=True)
                segment_data['vectorstore'].save_local(str(segment_path))
            
            # Générer un nom personnalisé si non fourni
            if not custom_name:
                sources = list(set([doc.metadata.get('source', 'Unknown')[:20] for doc in documents[:3]]))
                custom_name = f"VS_{len(sources)}docs_{datetime.now().strftime('%Y%m%d_%H%M')}"
            
            # Créer les métadonnées avec information de segmentation
            metadata = {
                'vectorstore_id': vectorstore_id,
                'custom_name': custom_name,
                'created_at': datetime.now().isoformat(),
                'embedding_model': embedding_model,
                'chunk_params': chunk_params,
                'document_count': len(documents),
                'vector_count': vectorstore.index.ntotal,
                'has_segments': True,
                'segments': {},
                'documents_info': []
            }
            
            # Ajouter les informations des documents et segments
            doc_sources = {}
            for doc in documents:
                source = doc.metadata.get('source', 'Unknown')
                doc_type = doc.metadata.get('type', 'unknown')
                doc_id = doc.metadata.get('doc_id', f"doc_{hashlib.md5(source.encode()).hexdigest()[:8]}")
                
                if source not in doc_sources:
                    doc_sources[source] = {
                        'source': source,
                        'type': doc_type,
                        'doc_id': doc_id,
                        'chunks': 0,
                        'total_size': 0
                    }
                
                doc_sources[source]['chunks'] += 1
                doc_sources[source]['total_size'] += len(doc.page_content)
            
            metadata['documents_info'] = list(doc_sources.values())
            
            # Ajouter les informations des segments
            for doc_id, segment_data in document_segments.items():
                metadata['segments'][doc_id] = {
                    'doc_id': doc_id,
                    'source': segment_data['source'],
                    'type': segment_data['type'],
                    'chunk_count': segment_data['chunk_count'],
                    'vector_count': segment_data['vectorstore'].index.ntotal
                }
            
            # Sauvegarder les métadonnées
            self.save_vectorstore_metadata(vectorstore_id, metadata)
            
            logger.info(f"Vector store segmenté sauvegardé: {vectorstore_path} avec {len(document_segments)} segments")
            return True
            
        except Exception as e:
            logger.error(f"Erreur sauvegarde vector store segmenté {vectorstore_id}: {str(e)}")
            return False
    
    def _create_document_segments(self, documents: List[LangchainDocument], 
                                embedding_model: str) -> Dict[str, Dict]:
        """Crée des segments de vector store pour chaque document"""
        # Regrouper les chunks par document source
        doc_groups = {}
        
        for doc in documents:
            doc_id = doc.metadata.get('doc_id')
            if not doc_id:
                # Générer un doc_id basé sur la source
                source = doc.metadata.get('source', 'Unknown')
                doc_id = f"doc_{hashlib.md5(source.encode()).hexdigest()[:8]}"
                doc.metadata['doc_id'] = doc_id
            
            if doc_id not in doc_groups:
                doc_groups[doc_id] = {
                    'documents': [],
                    'source': doc.metadata.get('source', 'Unknown'),
                    'type': doc.metadata.get('type', 'unknown')
                }
            
            doc_groups[doc_id]['documents'].append(doc)
        
        # Créer un vector store pour chaque groupe
        document_segments = {}
        embeddings = self._get_embeddings_instance(embedding_model)
        
        if embeddings:
            for doc_id, group_data in doc_groups.items():
                try:
                    # Créer le vector store pour ce document
                    segment_vectorstore = FAISS.from_documents(group_data['documents'], embeddings)
                    
                    document_segments[doc_id] = {
                        'vectorstore': segment_vectorstore,
                        'source': group_data['source'],
                        'type': group_data['type'],
                        'chunk_count': len(group_data['documents'])
                    }
                    
                except Exception as e:
                    logger.error(f"Erreur création segment pour {doc_id}: {str(e)}")
        
        return document_segments
    
    def _get_embeddings_instance(self, embedding_model: str):
        """Obtient une instance des embeddings pour la segmentation"""
        try:
            model_config = EMBEDDING_MODELS.get(embedding_model)
            if not model_config:
                return None
            
            # Vérifier si une clé API est requise
            if model_config["requires_api_key"]:
                if embedding_model == "OpenAI":
                    api_key = os.getenv("OPENAI_API_KEY")
                    if not api_key:
                        return None
            
            # Créer l'objet embeddings
            embedding_class = model_config["class"]
            params = model_config["params"].copy()
            
            embeddings = embedding_class(**params)
            return embeddings
            
        except Exception as e:
            logger.error(f"Erreur création embeddings pour segmentation: {str(e)}")
            return None
    
    def load_vectorstore(self, vectorstore_id: str, embeddings) -> Optional[FAISS]:
        """Charge un vector store complet depuis le répertoire"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            if not self.vectorstore_exists(vectorstore_id):
                logger.warning(f"Vector store {vectorstore_id} n'existe pas")
                return None
            
            # Charger le vector store
            vectorstore = FAISS.load_local(
                str(vectorstore_path), 
                embeddings, 
                allow_dangerous_deserialization=True
            )
            
            logger.info(f"Vector store chargé: {vectorstore_path}")
            return vectorstore
            
        except Exception as e:
            logger.error(f"Erreur chargement vector store {vectorstore_id}: {str(e)}")
            return None
    
    def load_document_segment(self, vectorstore_id: str, doc_id: str, embeddings) -> Optional[FAISS]:
        """Charge un segment spécifique d'un vector store"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            segment_path = vectorstore_path / "segments" / doc_id
            
            if not segment_path.exists() or not (segment_path / "index.faiss").exists():
                logger.warning(f"Segment {doc_id} n'existe pas dans {vectorstore_id}")
                return None
            
            # Charger le segment
            segment_vectorstore = FAISS.load_local(
                str(segment_path), 
                embeddings, 
                allow_dangerous_deserialization=True
            )
            
            logger.info(f"Segment chargé: {segment_path}")
            return segment_vectorstore
            
        except Exception as e:
            logger.error(f"Erreur chargement segment {doc_id}: {str(e)}")
            return None
    
    def load_multiple_segments(self, vectorstore_id: str, doc_ids: List[str], embeddings) -> Optional[FAISS]:
        """Charge et combine plusieurs segments de documents"""
        try:
            combined_vectorstore = None
            loaded_segments = []
            
            for doc_id in doc_ids:
                segment = self.load_document_segment(vectorstore_id, doc_id, embeddings)
                if segment:
                    if combined_vectorstore is None:
                        combined_vectorstore = segment
                    else:
                        combined_vectorstore.merge_from(segment)
                    loaded_segments.append(doc_id)
            
            if combined_vectorstore and loaded_segments:
                logger.info(f"Segments combinés: {', '.join(loaded_segments)}")
                return combined_vectorstore
            else:
                logger.warning("Aucun segment n'a pu être chargé")
                return None
                
        except Exception as e:
            logger.error(f"Erreur combinaison segments: {str(e)}")
            return None
    
    def list_vectorstores(self) -> List[Dict]:
        """Liste tous les vector stores disponibles avec leurs métadonnées"""
        vectorstores = []
        
        for path in self.base_dir.iterdir():
            if path.is_dir() and path.name.startswith("faiss_db_"):
                vectorstore_id = path.name.replace("faiss_db_", "")
                metadata = self.load_vectorstore_metadata(vectorstore_id)
                
                if metadata:
                    # Ajouter des informations sur la taille
                    total_size = sum(path.stat().st_size for path in path.rglob('*') if path.is_file())
                    metadata['storage_size_mb'] = round(total_size / (1024 * 1024), 2)
                    metadata['path'] = str(path)
                    
                    # Ajouter les segments disponibles
                    segments = self.get_document_segments(vectorstore_id)
                    metadata['available_segments'] = segments
                    
                    vectorstores.append(metadata)
        
        # Trier par date de création (plus récent d'abord)
        vectorstores.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        return vectorstores
    
    def delete_vectorstore(self, vectorstore_id: str) -> bool:
        """Supprime un vector store"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
                logger.info(f"Vector store supprimé: {vectorstore_id}")
                return True
            return False
        except Exception as e:
            logger.error(f"Erreur suppression vector store {vectorstore_id}: {str(e)}")
            return False
    
    def cleanup_old_vectorstores(self, max_age_days: int = 30, max_count: int = 10):
        """Nettoie les anciens vector stores"""
        vectorstores = self.list_vectorstores()
        current_time = datetime.now()
        deleted_count = 0
        
        # Supprimer les vector stores trop anciens
        for vs in vectorstores:
            created_at = datetime.fromisoformat(vs['created_at'])
            age_days = (current_time - created_at).days
            
            if age_days > max_age_days:
                self.delete_vectorstore(vs['vectorstore_id'])
                deleted_count += 1
        
        # Garder seulement les N plus récents
        if len(vectorstores) > max_count:
            for vs in vectorstores[max_count:]:
                self.delete_vectorstore(vs['vectorstore_id'])
                deleted_count += 1
        
        if deleted_count > 0:
            logger.info(f"Nettoyage: {deleted_count} vector stores supprimés")

class DocumentProcessor:
    """Classe pour traiter différents types de documents avec métadonnées"""
    
    def __init__(self):
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        # Dictionnaire pour stocker tous les documents par leur ID unique
        self.document_registry = {}
        
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def reset(self):
        """Remet à zéro le processor pour traiter de nouveaux documents"""
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        self.document_registry = {}
    
    def configure(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Configure les paramètres de traitement des documents"""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"Configuration mise à jour: chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def create_metadata(self, source: str, doc_type: str, size: int = 0, 
                       additional_info: Dict = None) -> Dict:
        """Crée les métadonnées pour un document avec ID unique"""
        doc_id = f"{doc_type}_{hashlib.md5(f'{source}_{size}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"

        
        metadata = {
            'doc_id': doc_id,
            'source': source,
            'type': doc_type,
            'size': size,
            'processed_at': datetime.now().isoformat(),
            'hash': hashlib.md5(source.encode()).hexdigest()[:8]
        }
        if additional_info:
            metadata.update(additional_info)
        return metadata
    
    def register_document(self, doc_metadata: Dict, content: str):
        """Enregistre un document dans le registre global"""
        doc_id = doc_metadata['doc_id']
        self.document_registry[doc_id] = {
            'metadata': doc_metadata,
            'content': content,
            'chunks': []
        }
    
    def process_pdf(self, pdf_file) -> List[LangchainDocument]:
        """Traite un fichier PDF avec métadonnées"""
        documents = []
        try:
            pdf_reader = PdfReader(pdf_file)
            full_text = ""
            doc_id = f"pdf_{hashlib.md5(f'{pdf_file.name}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    full_text += page_text + "\n\n"
                    
                    # Créer un document par page avec métadonnées
                    metadata = self.create_metadata(
                        source=pdf_file.name,
                        doc_type="pdf",
                        size=len(page_text),
                        additional_info={
                            'page_number': page_num + 1,
                            'doc_id': doc_id  # Même doc_id pour toutes les pages
                        }
                    )
                    metadata['doc_id'] = doc_id  # Forcer le même doc_id
                    
                    documents.append(LangchainDocument(
                        page_content=page_text,
                        metadata=metadata
                    ))
            
            # Enregistrer le document complet dans le registre
            if full_text.strip():
                complete_metadata = self.create_metadata(
                    source=pdf_file.name,
                    doc_type="pdf",
                    size=len(full_text),
                    additional_info={'total_pages': len(pdf_reader.pages)}
                )
                complete_metadata['doc_id'] = doc_id
                self.register_document(complete_metadata, full_text)
            
            self.stats['file_types']['pdf'] = self.stats['file_types'].get('pdf', 0) + 1
            logger.info(f"PDF traité avec succès: {pdf_file.name} ({len(documents)} pages)")
            
        except Exception as e:
            error_msg = f"Erreur PDF {pdf_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_docx(self, docx_file) -> List[LangchainDocument]:
        """Traite un fichier DOCX avec métadonnées"""
        documents = []
        try:
            # Essayer d'abord avec docx2txt
            text = docx2txt.process(docx_file)
            if not text.strip():
                # Fallback avec python-docx
                doc = Document(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if text.strip():
                doc_id = f"docx_{hashlib.md5(f'{docx_file.name}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"
                
                metadata = self.create_metadata(
                    source=docx_file.name,
                    doc_type="docx",
                    size=len(text),
                    additional_info={'word_count': len(text.split())}
                )
                metadata['doc_id'] = doc_id
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['docx'] = self.stats['file_types'].get('docx', 0) + 1
                logger.info(f"DOCX traité avec succès: {docx_file.name}")
            
        except Exception as e:
            error_msg = f"Erreur DOCX {docx_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_excel(self, excel_file) -> List[LangchainDocument]:
        """Traite un fichier Excel avec métadonnées"""
        documents = []
        try:
            # Déterminer le moteur selon l'extension
            engine = 'openpyxl' if excel_file.name.endswith('.xlsx') else 'xlrd'
            df_dict = pd.read_excel(excel_file, sheet_name=None, engine=engine)
            
            doc_id = f"excel_{hashlib.md5(f'{excel_file.name}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"
            
            for sheet_name, df in df_dict.items():
                if not df.empty:
                    # Créer un texte structuré pour chaque feuille
                    text_content = f"=== Feuille: {sheet_name} ===\n\n"
                    
                    # Ajouter les en-têtes
                    headers = list(df.columns)
                    text_content += "Colonnes: " + " | ".join(str(h) for h in headers) + "\n\n"
                    
                    # Ajouter les données
                    for index, row in df.iterrows():
                        row_data = []
                        for col in headers:
                            cell_value = row[col]
                            if pd.notna(cell_value):
                                row_data.append(f"{col}: {cell_value}")
                        
                        if row_data:
                            text_content += f"Ligne {index + 1} - " + " | ".join(row_data) + "\n"
                    
                    # Créer le document avec métadonnées
                    metadata = self.create_metadata(
                        source=excel_file.name,
                        doc_type="excel",
                        size=len(text_content),
                        additional_info={
                            'sheet_name': sheet_name,
                            'rows': len(df),
                            'columns': len(df.columns)
                        }
                    )
                    metadata['doc_id'] = doc_id  # Même doc_id pour toutes les feuilles
                    
                    documents.append(LangchainDocument(
                        page_content=text_content,
                        metadata=metadata
                    ))
            
            # Enregistrer le document complet dans le registre
            if documents:
                complete_text = "\n\n".join([doc.page_content for doc in documents])
                complete_metadata = self.create_metadata(
                    source=excel_file.name,
                    doc_type="excel",
                    size=len(complete_text),
                    additional_info={'total_sheets': len(df_dict)}
                )
                complete_metadata['doc_id'] = doc_id
                self.register_document(complete_metadata, complete_text)
            
            self.stats['file_types']['excel'] = self.stats['file_types'].get('excel', 0) + 1
            logger.info(f"Excel traité avec succès: {excel_file.name} ({len(documents)} feuilles)")
            
        except Exception as e:
            error_msg = f"Erreur Excel {excel_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_url(self, url: str) -> List[LangchainDocument]:
        """Traite une URL avec métadonnées"""
        documents = []
        try:
            if not self.is_valid_url(url):
                raise ValueError("URL invalide")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Parser le HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Supprimer les scripts et styles
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extraire le texte principal
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "Sans titre"
            
            # Chercher le contenu principal
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=['content', 'main', 'article'])
            
            if main_content:
                text = main_content.get_text()
            else:
                text = soup.get_text()
            
            # Nettoyer le texte
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if text.strip() and len(text) > 100:  # Ignorer le contenu trop court
                doc_id = f"url_{hashlib.md5(f'{url}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"
                
                metadata = self.create_metadata(
                    source=url,
                    doc_type="url",
                    size=len(text),
                    additional_info={
                        'title': title_text,
                        'domain': urllib.parse.urlparse(url).netloc,
                        'scraped_at': datetime.now().isoformat()
                    }
                )
                metadata['doc_id'] = doc_id
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['url'] = self.stats['file_types'].get('url', 0) + 1
                logger.info(f"URL traitée avec succès: {url}")
            else:
                raise ValueError("Contenu trop court ou vide")
                
        except Exception as e:
            error_msg = f"Erreur URL {url}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_text(self, text_content: str, source_name: str = "Texte personnalisé") -> List[LangchainDocument]:
        """Traite du texte brut avec métadonnées"""
        documents = []
        try:
            if text_content.strip():
                doc_id = f"text_{hashlib.md5(f'{source_name}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"
                
                metadata = self.create_metadata(
                    source=source_name,
                    doc_type="text",
                    size=len(text_content),
                    additional_info={
                        'word_count': len(text_content.split()),
                        'line_count': len(text_content.splitlines())
                    }
                )
                metadata['doc_id'] = doc_id
                
                # Enregistrer dans le registre
                self.register_document(metadata, text_content)
                
                documents.append(LangchainDocument(
                    page_content=text_content,
                    metadata=metadata
                ))
                
                self.stats['file_types']['text'] = self.stats['file_types'].get('text', 0) + 1
                logger.info(f"Texte traité avec succès: {source_name}")
                
        except Exception as e:
            error_msg = f"Erreur texte {source_name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def is_valid_url(self, url: str) -> bool:
        """Vérifie si une URL est valide"""
        try:
            result = urllib.parse.urlparse(url)
            return all([result.scheme, result.netloc])
        except Exception:
            return False
    
    def process_files(self, files) -> List[LangchainDocument]:
        """Traite une liste de fichiers uploadés"""
        all_documents = []
        start_time = time.time()
        
        if files:
            self.stats['total_docs'] = len(files)
            
            for file in files:
                try:
                    file_extension = file.name.lower().split('.')[-1]
                    
                    if file_extension == 'pdf':
                        docs = self.process_pdf(file)
                    elif file_extension in ['docx', 'doc']:
                        docs = self.process_docx(file)
                    elif file_extension in ['xlsx', 'xls']:  # Ajout de 'xls'
                        docs = self.process_excel(file)
                    else:
                        # Traiter comme texte
                        try:
                            content = file.read().decode('utf-8')
                            docs = self.process_text(content, file.name)
                        except UnicodeDecodeError:
                            # Essayer avec d'autres encodages
                            file.seek(0)
                            try:
                                content = file.read().decode('latin-1')
                                docs = self.process_text(content, file.name)
                            except Exception:
                                docs = []
                                self.errors.append(f"Impossible de décoder le fichier: {file.name}")
                    
                    if docs:
                        all_documents.extend(docs)
                        self.stats['successful_docs'] += 1
                    else:
                        self.stats['failed_docs'] += 1
                        
                except Exception as e:
                    self.stats['failed_docs'] += 1
                    error_msg = f"Erreur traitement {file.name}: {str(e)}"
                    self.errors.append(error_msg)
                    logger.error(error_msg)
    
        self.stats['processing_time'] = time.time() - start_time
        self.processed_docs = all_documents
        
        return all_documents
    
    def get_statistics(self) -> Dict:
        """Retourne les statistiques de traitement"""
        return self.stats.copy()
    
    def get_errors(self) -> List[str]:
        """Retourne la liste des erreurs"""
        return self.errors.copy()
    
    def get_document_info(self, doc_id: str) -> Optional[Dict]:
        """Récupère les informations d'un document par son ID"""
        return self.document_registry.get(doc_id)
    
    def list_processed_documents(self) -> List[Dict]:
        """Liste tous les documents traités avec leurs métadonnées"""
        docs_info = []
        for doc_id, doc_info in self.document_registry.items():
            metadata = doc_info['metadata'].copy()
            metadata['content_preview'] = doc_info['content'][:200] + "..." if len(doc_info['content']) > 200 else doc_info['content']
            docs_info.append(metadata)
        
        return docs_info

def get_embeddings(embedding_model: str):
    """Crée l'objet embeddings selon le modèle sélectionné"""
    try:
        model_config = EMBEDDING_MODELS.get(embedding_model)
        if not model_config:
            st.error(f"Modèle d'embedding non supporté: {embedding_model}")
            return None
        
        # Vérifier si une clé API est requise
        if model_config["requires_api_key"]:
            if embedding_model == "OpenAI":
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                    st.error("Clé API OpenAI manquante dans les variables d'environnement")
                    return None
        
        # Créer l'objet embeddings
        embedding_class = model_config["class"]
        params = model_config["params"].copy()
        
        with st.spinner(f"Initialisation du modèle d'embedding {embedding_model}..."):
            embeddings = embedding_class(**params)
        
        return embeddings
        
    except Exception as e:
        st.error(f"Erreur lors de la création des embeddings {embedding_model}: {str(e)}")
        logger.error(f"Erreur embeddings {embedding_model}: {str(e)}")
        return None

def get_text_chunks(documents: List[LangchainDocument], chunk_size: int = 1000, 
                   chunk_overlap: int = 200) -> List[LangchainDocument]:
    """Divise les documents en chunks avec métadonnées préservées"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    all_chunks = []
    for doc in documents:
        # Diviser le document en chunks
        chunks = text_splitter.split_text(doc.page_content)
        
        # Créer des documents LangChain avec métadonnées préservées
        for i, chunk_text in enumerate(chunks):
            # Copier les métadonnées originales
            chunk_metadata = doc.metadata.copy()
            
            # Ajouter des informations sur le chunk
            chunk_metadata.update({
                'chunk_index': i,
                'chunk_id': f"{chunk_metadata.get('doc_id', 'unknown')}_{i}",
                'chunk_size': len(chunk_text),
                'total_chunks': len(chunks)
            })
            
            chunk_doc = LangchainDocument(
                page_content=chunk_text,
                metadata=chunk_metadata
            )
            all_chunks.append(chunk_doc)
    
    logger.info(f"Documents divisés en {len(all_chunks)} chunks")
    return all_chunks

def create_vectorstore(text_chunks: List[LangchainDocument], embeddings, 
                      vs_manager: SegmentedVectorStoreManager,
                      embedding_model: str, chunk_params: Dict,
                      custom_name: str = None) -> Tuple[Optional[FAISS], Optional[str]]:
    """Crée et sauvegarde un vector store avec gestion de la segmentation"""
    try:
        if not text_chunks:
            st.error("Aucun chunk de texte disponible pour créer le vector store")
            return None, None
        
        # Générer un ID unique pour ce vector store
        vectorstore_id = vs_manager.generate_vectorstore_id(
            text_chunks, embedding_model, chunk_params
        )
        
        # Vérifier si ce vector store existe déjà
        if vs_manager.vectorstore_exists(vectorstore_id):
            st.info("Un vector store identique existe déjà, chargement...")
            vectorstore = vs_manager.load_vectorstore(vectorstore_id, embeddings)
            return vectorstore, vectorstore_id
        
        # Créer le vector store
        with st.spinner("Création du vector store en cours..."):
            vectorstore = FAISS.from_documents(text_chunks, embeddings)
        
        # Sauvegarder avec segmentation
        success = vs_manager.save_segmented_vectorstore(
            vectorstore, vectorstore_id, text_chunks, 
            embedding_model, chunk_params, custom_name
        )
        
        if success:
            st.success(f"Vector store créé et sauvegardé: {vectorstore_id}")
            return vectorstore, vectorstore_id
        else:
            st.warning("Vector store créé mais non sauvegardé")
            return vectorstore, vectorstore_id
            
    except Exception as e:
        st.error(f"Erreur lors de la création du vector store: {str(e)}")
        logger.error(f"Erreur création vector store: {str(e)}")
        return None, None

def get_llm(model_name: str, temperature: float = 0.3):
    """Crée l'instance du modèle LLM"""
    try:
        model_config = LLM_MODELS.get(model_name)
        if not model_config:
            st.error(f"Modèle LLM non supporté: {model_name}")
            return None
        
        api_key = os.getenv(model_config["api_key_env"])
        if not api_key:
            st.error(f"Clé API manquante pour {model_name}: {model_config['api_key_env']}")
            return None
        
        llm = ChatGoogleGenerativeAI(
            model=model_config["model_name"],
            temperature=temperature,
            google_api_key=api_key
        )
        
        return llm
        
    except Exception as e:
        st.error(f"Erreur lors de la création du LLM {model_name}: {str(e)}")
        logger.error(f"Erreur LLM {model_name}: {str(e)}")
        return None

def get_conversation_chain(vectorstore, llm):
    """Crée la chaîne de conversation avec mémoire"""
    try:
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 6}
            ),
            memory=memory,
            return_source_documents=True,
            verbose=True
        )
        
        return conversation_chain
        
    except Exception as e:
        st.error(f"Erreur lors de la création de la chaîne de conversation: {str(e)}")
        logger.error(f"Erreur conversation chain: {str(e)}")
        return None
def handle_user_input(user_question):
    """Gère les questions de l'utilisateur et affiche les réponses"""
    if user_question and st.session_state.conversation:
        try:
            with st.spinner("Génération de la réponse..."):
                # Utiliser la clé 'question' pour ConversationalRetrievalChain
                response = st.session_state.conversation({
                    "question": user_question
                })
                
                # Extraire la réponse et les documents sources
                answer = response.get('answer', 'Désolé, je n\'ai pas pu générer une réponse.')
                source_documents = response.get('source_documents', [])
                
                # Sauvegarder dans l'historique
                if 'chat_history' not in st.session_state:
                    st.session_state.chat_history = []
                
                st.session_state.chat_history.append({
                    'question': user_question,
                    'answer': answer,
                    'sources': source_documents,
                    'timestamp': datetime.now().isoformat()
                })
                
        except Exception as e:
            st.error(f"Erreur lors du traitement de la question: {str(e)}")
            logger.error(f"Erreur handle_user_input: {str(e)}")

def display_chat_history():
    """Affiche l'historique des conversations"""
    if 'chat_history' in st.session_state and st.session_state.chat_history:
        st.subheader("💬 Historique des conversations")
        
        for i, exchange in enumerate(reversed(st.session_state.chat_history)):
            with st.container():
                # Question de l'utilisateur
                st.write(user_template.replace("{{MSG}}", exchange['question']), 
                        unsafe_allow_html=True)
                
                # Réponse du bot
                st.write(bot_template.replace("{{MSG}}", exchange['answer']), 
                        unsafe_allow_html=True)
                
                # Sources utilisées
                if exchange.get('sources'):
                    with st.expander(f"📚 Sources utilisées ({len(exchange['sources'])})"):
                        for j, doc in enumerate(exchange['sources']):
                            source = doc.metadata.get('source', 'Source inconnue')
                            doc_type = doc.metadata.get('type', 'unknown')
                            chunk_info = doc.metadata.get('chunk_index', 'N/A')
                            
                            st.write(f"**Source {j+1}:** {source} ({doc_type})")
                            if chunk_info != 'N/A':
                                st.write(f"*Chunk {chunk_info}*")
                            
                            # Afficher un extrait du contenu
                            content_preview = doc.page_content[:300]
                            if len(doc.page_content) > 300:
                                content_preview += "..."
                            st.write(f"```\n{content_preview}\n```")
                            st.write("---")
                
                st.write("---")

def display_vectorstore_manager():
    """Affiche le gestionnaire de vector stores"""
    st.subheader("🗂️ Gestionnaire de Vector Stores")
    
    # Initialiser le gestionnaire si nécessaire
    if 'vs_manager' not in st.session_state:
        st.session_state.vs_manager = SegmentedVectorStoreManager()
    
    vs_manager = st.session_state.vs_manager
    
    # Onglets pour les différentes fonctions
    tab1, tab2, tab3 = st.tabs(["📋 Vector Stores", "🔍 Segments", "🗑️ Gestion"])
    
    with tab1:
        # Liste des vector stores disponibles
        vectorstores = vs_manager.list_vectorstores()
        
        if vectorstores:
            st.write(f"**{len(vectorstores)} vector store(s) disponible(s):**")
            
            for vs in vectorstores:
                with st.expander(f"📁 {vs.get('custom_name', vs['vectorstore_id'])} - {vs.get('storage_size_mb', 0):.1f} MB"):
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.write(f"**ID:** {vs['vectorstore_id']}")
                        st.write(f"**Créé le:** {vs.get('created_at', 'N/A')}")
                        st.write(f"**Modèle d'embedding:** {vs.get('embedding_model', 'N/A')}")
                        st.write(f"**Documents:** {vs.get('document_count', 0)}")
                        st.write(f"**Vecteurs:** {vs.get('vector_count', 0)}")
                        st.write(f"**Segments:** {len(vs.get('available_segments', []))}")
                        
                        # Informations sur les documents
                        if vs.get('documents_info'):
                            st.write("**Sources:**")
                            for doc_info in vs['documents_info']:
                                st.write(f"- {doc_info['source']} ({doc_info['type']}) - {doc_info['chunks']} chunks")
                    
                    with col2:
                        # Bouton pour charger ce vector store
                        if st.button(f"📥 Charger", key=f"load_{vs['vectorstore_id']}"):
                            try:
                                # Récupérer le modèle d'embedding utilisé
                                embedding_model = vs.get('embedding_model', 'HuggingFace Sentence Transformers')
                                embeddings = get_embeddings(embedding_model)
                                
                                if embeddings:
                                    loaded_vs = vs_manager.load_vectorstore(vs['vectorstore_id'], embeddings)
                                    if loaded_vs:
                                        st.session_state.vectorstore = loaded_vs
                                        st.session_state.current_vectorstore_id = vs['vectorstore_id']
                                        st.success(f"Vector store chargé: {vs.get('custom_name', vs['vectorstore_id'])}")
                                        st.rerun()
                                    else:
                                        st.error("Erreur lors du chargement du vector store")
                                else:
                                    st.error("Impossible de créer les embeddings")
                            except Exception as e:
                                st.error(f"Erreur: {str(e)}")
        else:
            st.info("Aucun vector store disponible. Créez-en un en traitant des documents.")
    
    with tab2:
        # Gestion des segments
        if 'current_vectorstore_id' in st.session_state:
            current_vs_id = st.session_state.current_vectorstore_id
            segments = vs_manager.get_document_segments(current_vs_id)
            
            if segments:
                st.write(f"**Segments disponibles pour le vector store actuel:**")
                
                selected_segments = []
                for segment in segments:
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        is_selected = st.checkbox(
                            f"📄 {segment['source']} ({segment['type']}) - {segment['chunk_count']} chunks",
                            key=f"segment_{segment['doc_id']}"
                        )
                        if is_selected:
                            selected_segments.append(segment['doc_id'])
                    
                    with col2:
                        st.write(f"{segment['vector_count']} vecteurs")
                
                # Bouton pour charger les segments sélectionnés
                if selected_segments:
                    if st.button("🔄 Charger segments sélectionnés"):
                        try:
                            metadata = vs_manager.load_vectorstore_metadata(current_vs_id)
                            embedding_model = metadata.get('embedding_model', 'HuggingFace Sentence Transformers')
                            embeddings = get_embeddings(embedding_model)
                            if embeddings:
                                combined_vs = vs_manager.load_multiple_segments(
                                    current_vs_id, selected_segments, embeddings
                                )
                                if combined_vs:
                                    st.session_state.vectorstore = combined_vs
                                    st.success(f"Segments combinés chargés: {len(selected_segments)} documents")
                                    # Créer la chaîne de conversation avec le vectorstore des segments
                                    st.session_state.conversation = get_conversation_chain(combined_vs, get_llm(st.session_state.selected_llm_model))
                                    st.rerun()
                                else:
                                    st.error("Erreur lors du chargement des segments")
                            else:
                                st.error("Impossible de créer les embeddings")
                        except Exception as e:
                            st.error(f"Erreur: {str(e)}")
            else:
                st.info("Aucun segment disponible pour le vector store actuel.")
        else:
            st.info("Aucun vector store actuel. Chargez ou créez un vector store d'abord.")
    
    with tab3:
        # Nettoyage et suppression
        st.write("**🧹 Nettoyage automatique:**")
        
        col1, col2 = st.columns(2)
        with col1:
            max_age = st.number_input("Âge max (jours)", min_value=1, max_value=365, value=30)
        with col2:
            max_count = st.number_input("Nombre max", min_value=1, max_value=100, value=10)
        
        if st.button("🧹 Nettoyer anciens vector stores"):
            try:
                vs_manager.cleanup_old_vectorstores(max_age, max_count)
                st.success("Nettoyage effectué")
                st.rerun()
            except Exception as e:
                st.error(f"Erreur nettoyage: {str(e)}")
        
        st.write("**🗑️ Suppression manuelle:**")
        vectorstores = vs_manager.list_vectorstores()
        
        if vectorstores:
            vs_to_delete = st.selectbox(
                "Sélectionner un vector store à supprimer:",
                options=[None] + [vs['vectorstore_id'] for vs in vectorstores],
                format_func=lambda x: "Choisir..." if x is None else next(
                    (vs.get('custom_name', vs['vectorstore_id']) for vs in vectorstores if vs['vectorstore_id'] == x),
                    x
                )
            )
            
            if vs_to_delete:
                if st.button("🗑️ Supprimer définitivement", type="secondary"):
                    try:
                        success = vs_manager.delete_vectorstore(vs_to_delete)
                        if success:
                            st.success("Vector store supprimé")
                            # Nettoyer la session si c'était le vector store actuel
                            if st.session_state.get('current_vectorstore_id') == vs_to_delete:
                                if 'vectorstore' in st.session_state:
                                    del st.session_state.vectorstore
                                if 'current_vectorstore_id' in st.session_state:
                                    del st.session_state.current_vectorstore_id
                            st.rerun()
                        else:
                            st.error("Erreur lors de la suppression")
                    except Exception as e:
                        st.error(f"Erreur: {str(e)}")

def display_document_processor_stats():
    """Affiche les statistiques du processeur de documents"""
    if 'doc_processor' in st.session_state:
        processor = st.session_state.doc_processor
        stats = processor.get_statistics()
        errors = processor.get_errors()
        
        if stats['total_docs'] > 0:
            st.subheader("📊 Statistiques de traitement")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Documents traités", stats['successful_docs'])
            with col2:
                st.metric("Échecs", stats['failed_docs'])
            with col3:
                st.metric("Temps de traitement", f"{stats['processing_time']:.2f}s")
            with col4:
                success_rate = (stats['successful_docs'] / stats['total_docs'] * 100) if stats['total_docs'] > 0 else 0
                st.metric("Taux de succès", f"{success_rate:.1f}%")
            
            # Types de fichiers traités
            if stats['file_types']:
                st.write("**Types de fichiers:**")
                for file_type, count in stats['file_types'].items():
                    st.write(f"- {file_type.upper()}: {count}")
            
            # Erreurs
            if errors:
                with st.expander(f"⚠️ Erreurs ({len(errors)})"):
                    for error in errors:
                        st.error(error)
            
            # Informations sur les documents traités
            docs_info = processor.list_processed_documents()
            if docs_info:
                with st.expander(f"📄 Documents traités ({len(docs_info)})"):
                    for doc_info in docs_info:
                        st.write(f"**{doc_info['source']}** ({doc_info['type']})")
                        st.write(f"- Taille: {doc_info['size']} caractères")
                        st.write(f"- ID: {doc_info['doc_id']}")
                        if 'word_count' in doc_info:
                            st.write(f"- Mots: {doc_info['word_count']}")
                        st.write(f"- Aperçu: {doc_info.get('content_preview', 'N/A')}")
                        st.write("---")

def main():
    """Fonction principale de l'application Streamlit"""
    # Configuration de la page
    st.set_page_config(
        page_title="Chat Multi-Documents avec Vector Stores",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Charger les variables d'environnement
    load_dotenv()
    
    # CSS personnalisé
    st.write(css, unsafe_allow_html=True)
    
    # Titre principal
    st.title("🤖 Chat Multi-Documents avec Vector Stores Segmentés")
    st.markdown("---")
    
    # Initialiser le gestionnaire de vector stores
    if 'vs_manager' not in st.session_state:
        st.session_state.vs_manager = SegmentedVectorStoreManager()
    
    # Initialiser le processeur de documents
    if 'doc_processor' not in st.session_state:
        st.session_state.doc_processor = DocumentProcessor()
    
    # Sidebar pour la configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Section modèles
        st.subheader("🧠 Modèles")
        
        # Sélection du modèle d'embedding
        embedding_model = st.selectbox(
            "Modèle d'Embedding:",
            options=list(EMBEDDING_MODELS.keys()),
            index=2,  # HuggingFace
        )
        st.session_state.selected_embedding_model = embedding_model
        
        # Sélection du modèle LLM
        llm_model = st.selectbox(
            "Modèle LLM:",
            options=list(LLM_MODELS.keys()),
            index=0
        )
        st.session_state.selected_llm_model = llm_model
        
        # Configuration avancée
        with st.expander("🔧 Paramètres avancés"):
            chunk_size = st.slider("Taille des chunks", 500, 2000, 1000, 100)
            chunk_overlap = st.slider("Chevauchement", 50, 300, 200, 50)
            k_documents = st.slider("Nombre de documents à récupérer", 1, 10, 4)
            temperature = st.slider("Température LLM", 0.0, 1.0, 0.3, 0.1)
            
            st.session_state.chunk_size = chunk_size
            st.session_state.chunk_overlap = chunk_overlap
            st.session_state.k_documents = k_documents
            st.session_state.temperature = temperature
        
        # Informations sur le vector store actuel
        if 'current_vectorstore_id' in st.session_state:
            st.success(f"✅ Vector store actuel: {st.session_state.current_vectorstore_id[:8]}...")
            if st.button("🔄 Recharger conversation"):
                create_conversation_chain()
        else:
            st.info("ℹ️ Aucun vector store chargé")
        
        # Bouton pour effacer l'historique
        if st.button("🗑️ Effacer l'historique"):
            if 'chat_history' in st.session_state:
                st.session_state.chat_history = []
            st.success("Historique effacé")
            st.rerun()
    
    # Interface principale avec onglets
    tab1, tab2, tab3, tab4 = st.tabs([
        "💬 Chat", 
        "📁 Documents", 
        "🗂️ Vector Stores", 
        "📊 Statistiques"
    ])
    
    with tab1:
        # Interface de chat
        st.subheader("💬 Conversation avec vos documents")
        user_question = st.text_input(
            "Posez votre question:",
            placeholder="Que voulez-vous savoir sur vos documents ?",
            key="user_input"
        )
        send_button = st.button("📤 Envoyer", type="primary")

        # Seul le bouton déclenche la génération
        if send_button and user_question:
            if 'vectorstore' not in st.session_state:
                st.warning("⚠️ Veuillez d'abord charger ou créer un vector store.")
            else:
                if 'conversation' not in st.session_state:
                    create_conversation_chain()
                handle_user_input(user_question)
                # Optionnel : st.session_state.user_input = "" (mais évite st.rerun ici)
        display_chat_history()
    
    with tab2:
        # Interface de gestion des documents
        st.subheader("📁 Traitement des documents")
        
        # Upload de fichiers
        uploaded_files = st.file_uploader(
            "Choisissez vos fichiers",
            type=['pdf', 'txt', 'docx', 'xlsx', 'xls', 'csv', 'pptx'],  # Ajout de 'xls'
            accept_multiple_files=True,
            help="Formats supportés: PDF, TXT, DOCX, XLSX, XLS, CSV, PPTX"
        )
        
        # Ajout du champ pour les URLs
        urls_input = st.text_area(
            "URLs à traiter (une par ligne):",
            placeholder="https://example.com\nhttps://another-site.com",
            height=100
        )
        
        # Options de traitement
        col1, col2 = st.columns(2)
        with col1:
            custom_name = st.text_input(
                "Nom personnalisé du vector store:",
                placeholder="Optionnel - sera généré automatiquement si vide"
            )
        
        with col2:
            processing_mode = st.selectbox(
                "Mode de traitement:",
                options=[
                    "Nouveau vector store",
                    "Ajouter au vector store actuel"
                ]
            )
        
        # Bouton de traitement
        if st.button("🚀 Traiter les documents", type="primary"):
            if not uploaded_files and not urls_input.strip():
                st.warning("⚠️ Veuillez sélectionner au moins un fichier ou entrer une URL.")
            else:
                try:
                    with st.spinner("Traitement en cours..."):
                        # Obtenir les embeddings
                        embeddings = get_embeddings(st.session_state.selected_embedding_model)
                        if not embeddings:
                            st.error("Impossible de créer les embeddings")
                            return
                        
                        # Traitement des documents
                        documents = []
                        
                        # Traiter les fichiers uploadés
                        if uploaded_files:
                            file_docs = st.session_state.doc_processor.process_files(uploaded_files)
                            documents.extend(file_docs)
                        
                        # Traiter les URLs
                        if urls_input.strip():
                            urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
                            for url in urls:
                                if st.session_state.doc_processor.is_valid_url(url):
                                    url_docs = st.session_state.doc_processor.process_url(url)
                                    if url_docs:
                                        documents.extend(url_docs)
                                        st.success(f"✅ URL traitée: {url}")
                                else:
                                    st.warning(f"⚠️ URL invalide ignorée: {url}")
                        
                        if not documents:
                            st.error("❌ Aucun document n'a pu être traité.")
                            return
                        
                        # Créer ou mettre à jour le vector store
                        chunk_params = {
                            'chunk_size': st.session_state.get('chunk_size', 1000),
                            'chunk_overlap': st.session_state.get('chunk_overlap', 200)
                        }
                        
                        # Diviser les documents en chunks
                        text_chunks = get_text_chunks(
                            documents,
                            chunk_size=chunk_params['chunk_size'],
                            chunk_overlap=chunk_params['chunk_overlap']
                        )
                        
                        if processing_mode == "Nouveau vector store":
                            # Créer un nouveau vector store
                            vectorstore, vectorstore_id = create_vectorstore(
                                text_chunks=text_chunks,
                                embeddings=embeddings,
                                vs_manager=st.session_state.vs_manager,
                                embedding_model=st.session_state.selected_embedding_model,
                                chunk_params=chunk_params,
                                custom_name=custom_name
                            )
                            
                            if vectorstore and vectorstore_id:
                                st.session_state.vectorstore = vectorstore
                                st.session_state.current_vectorstore_id = vectorstore_id
                                st.success(f"✅ Nouveau vector store créé avec les documents et URLs: {vectorstore_id}")
                                create_conversation_chain()
                                st.rerun()
                        else:
                            # Ajouter au vector store existant
                            if 'current_vectorstore_id' not in st.session_state:
                                st.error("❌ Aucun vector store actuel sélectionné")
                                return
                            
                            # ... code pour ajouter au vector store existant ...
                            st.error("⚠️ Fonctionnalité en cours de développement")
                        
                        # Afficher les statistiques
                        display_document_processor_stats()
                        
                except Exception as e:
                    st.error(f"❌ Erreur lors du traitement: {str(e)}")
                    logger.error(f"Erreur traitement documents et URLs: {str(e)}")
    
    with tab3:
        # Gestionnaire de vector stores
        display_vectorstore_manager()
    
    with tab4:
        # Statistiques et informations
        st.subheader("📊 Statistiques globales")
        
        # Informations sur le système
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**🔧 Configuration actuelle:**")
            st.write(f"- Modèle d'embedding: {st.session_state.get('selected_embedding_model', 'Non défini')}")
            st.write(f"- Modèle LLM: {st.session_state.get('selected_llm_model', 'Non défini')}")
            st.write(f"- Taille des chunks: {st.session_state.get('chunk_size', 1000)}")
            st.write(f"- Chevauchement: {st.session_state.get('chunk_overlap', 200)}")
            st.write(f"- Documents récupérés: {st.session_state.get('k_documents', 4)}")
            st.write(f"- Température: {st.session_state.get('temperature', 0.3)}")
        
        with col2:
            st.write("**📈 Statistiques de session:**")
            chat_count = len(st.session_state.get('chat_history', []))
            st.write(f"- Messages échangés: {chat_count}")
            
            if 'vectorstore' in st.session_state and hasattr(st.session_state.vectorstore, 'index'):
                try:
                    vector_count = st.session_state.vectorstore.index.ntotal
                    st.write(f"- Vecteurs dans le store: {vector_count}")
                except:
                    st.write("- Vecteurs dans le store: N/A")
            
            if 'vs_manager' in st.session_state:
                vs_count = len(st.session_state.vs_manager.list_vectorstores())
                st.write(f"- Vector stores disponibles: {vs_count}")
        
        # Informations détaillées sur le vector store actuel
        if 'current_vectorstore_id' in st.session_state:
            st.write("**🗂️ Vector store actuel:**")
            current_vs_id = st.session_state.current_vectorstore_id
            vs_manager = st.session_state.vs_manager
            
            try:
                metadata = vs_manager.load_vectorstore_metadata(current_vs_id)
                if metadata:
                    st.write(f"- ID: {current_vs_id}")
                    st.write(f"- Nom: {metadata.get('custom_name', 'N/A')}")
                    st.write(f"- Créé le: {metadata.get('created_at', 'N/A')}")
                    st.write(f"- Taille: {metadata.get('storage_size_mb', 0):.2f} MB")
                    st.write(f"- Documents: {metadata.get('document_count', 0)}")
                    
                    # Graphique des types de documents
                    if metadata.get('documents_info'):
                        doc_types = {}
                        for doc_info in metadata['documents_info']:
                            doc_type = doc_info['type']
                            doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
                        
                        if doc_types:
                            st.write("**📊 Répartition par type:**")
                            for doc_type, count in doc_types.items():
                                percentage = (count / len(metadata['documents_info'])) * 100
                                st.write(f"- {doc_type.upper()}: {count} ({percentage:.1f}%)")
            except Exception as e:
                st.error(f"Erreur lors de la récupération des métadonnées: {str(e)}")
        
        # Actions de maintenance
        st.subheader("🔧 Maintenance")
        
        if st.button("🔄 Réinitialiser session"):
            # Conserver seulement les éléments essentiels
            keys_to_keep = ['vs_manager']
            session_backup = {k: v for k, v in st.session_state.items() if k in keys_to_keep}
            st.session_state.clear()
            st.session_state.update(session_backup)
            st.success("Session réinitialisée")
            st.rerun()
        
        with col3:
            if st.button("📥 Exporter configuration"):
                config = {
                    'embedding_model': st.session_state.get('selected_embedding_model'),
                    'llm_model': st.session_state.get('selected_llm_model'),
                    'chunk_size': st.session_state.get('chunk_size', 1000),
                    'chunk_overlap': st.session_state.get('chunk_overlap', 200),
                    'k_documents': st.session_state.get('k_documents', 4),
                    'temperature': st.session_state.get('temperature', 0.3)
                }
                
                import json
                config_json = json.dumps(config, indent=2)
                
                
                st.download_button(
                    label="📥 Télécharger config.json",
                    data=config_json,
                    file_name="chat_config.json",
                    mime="application/json"
                )

def create_conversation_chain():
    """Crée la chaîne de conversation RAG"""
    if 'vectorstore' not in st.session_state:
        logger.warning("Aucun vectorstore disponible pour créer la chaîne de conversation")
        return False
    
    try:
        # Obtenir le modèle LLM
        llm_model_name = st.session_state.get('selected_llm_model')
        llm = get_llm(llm_model_name)
        
        if not llm:
            st.error("Impossible de créer le modèle LLM")
            return False
            
        # Créer la mémoire
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        # Créer la chaîne de conversation
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=st.session_state.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": st.session_state.get('k_documents', 4)}
            ),
            memory=memory,
            return_source_documents=True,
            verbose=True
        )
        
        st.session_state.conversation = conversation_chain
        logger.info("Chaîne de conversation créée avec succès")
        return True
        
    except Exception as e:
        logger.error(f"Erreur création chaîne conversation: {str(e)}")
        st.error(f"Erreur lors de la création de la conversation: {str(e)}")
        return False

# CSS pour l'interface
# css = """
# <style>
# .stTextInput > div > div > input {
#     background-color: #f0f2f6;
# }

# .stSelectbox > div > div > select {
#     background-color: #f0f2f6;
# }

# .stButton > button {
#     background-color: #0083b8;
#     color: white;
#     border-radius: 5px;
#     border: none;
#     padding: 0.5rem 1rem;
#     font-weight: bold;
# }

# .stButton > button:hover {
#     background-color: #006c96;
# }

# .chat-message {
#     padding: 1rem;
#     border-radius: 10px;
#     margin: 1rem 0;
#     display: flex;
#     align-items: flex-start;
# }

# .chat-message.user {
#     background-color: #dcf8c6;
#     flex-direction: row-reverse;
# }

# .chat-message.bot {
#     background-color: #f1f1f1;
# }

# .chat-message .avatar {
#     width: 40px;
#     height: 40px;
#     border-radius: 50%;
#     object-fit: cover;
#     margin: 0 10px;
# }

# .chat-message .message {
#     flex-grow: 1;
#     padding: 0 10px;
# }

# .source-doc {
#     background-color: #f8f9fa;
#     border-left: 4px solid #0083b8;
#     padding: 10px;
#     margin: 5px 0;
#     border-radius: 5px;
# }

# .metric-container {
#     background-color: #f0f2f6;
#     padding: 1rem;
#     border-radius: 10px;
#     text-align: center;
# }

# .error-message {
#     background-color: #ffebee;
#     color: #c62828;
#     padding: 10px;
#     border-radius: 5px;
#     border-left: 4px solid #c62828;
#     margin: 10px 0;
# }

# .success-message {
#     background-color: #e8f5e8;
#     color: #2e7d32;
#     padding: 10px;
#     border-radius: 5px;
#     border-left: 4px solid #2e7d32;
#     margin: 10px 0;
# }

# .info-message {
#     background-color: #e3f2fd;
#     color: #1565c0;
#     padding: 10px;
#     border-radius: 5px;
#     border-left: 4px solid #1565c0;
#     margin: 10px 0;
# }
# </style>
# """

# Templates HTML pour les messages
# user_template = """
# <div class="chat-message user">
#     <div class="message">{{MSG}}</div>
#     <div class="avatar">👤</div>
# </div>
# """

# bot_template = """
# <div class="chat-message bot">
#     <div class="avatar">🤖</div>
#     <div class="message">{{MSG}}</div>
# </div>
# """

if __name__ == "__main__":
    main()
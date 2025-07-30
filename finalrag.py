import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings, HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.document_loaders import WebBaseLoader
import requests
from bs4 import BeautifulSoup
import urllib.parse
import pandas as pd
import openpyxl
import xlrd
from pydantic import BaseModel
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument
from htmlTemplates import css, bot_template, user_template

import os
import warnings
import hashlib
import json
import shutil
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import logging
from pathlib import Path
import time
import numpy as np
import pickle

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Supprimer les avertissements
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Répertoires de stockage
VECTORSTORE_BASE_DIR = "vectorstores"
CACHE_DIR = "cache"

# Configuration des modèles
EMBEDDING_MODELS = {
    "OpenAI": {
        "class": OpenAIEmbeddings,
        "params": {},
        "requires_api_key": True
    },
    "HuggingFace Instructor": {
        "class": HuggingFaceInstructEmbeddings,
        "params": {"model_name": "hku-nlp/instructor-xl"},
        "requires_api_key": False
    },
    "HuggingFace Sentence Transformers": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/all-MiniLM-L6-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    },
    "HuggingFace BGE Small": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "BAAI/bge-small-en-v1.5",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    },
    "HuggingFace Multilingual": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False
    }
}

LLM_MODELS = {
    "Gemini 2.0 Flash": {
        "model_name": "gemini-2.0-flash-exp",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle puissant et rapide"
    },
    "Gemini Pro 1.5": {
        "model_name": "gemini-1.5-pro-latest",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini 1.5 Pro - Modèle le plus avancé"
    },
    "Gemini Pro": {
        "model_name": "gemini-pro",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle standard performant"
    }
}

class VectorStoreManager:
    """Gestionnaire des vector stores FAISS avec sauvegarde en répertoire - Version corrigée"""
    
    def __init__(self, base_dir: str = VECTORSTORE_BASE_DIR):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(exist_ok=True)
        self.cache_dir = Path(CACHE_DIR)
        self.cache_dir.mkdir(exist_ok=True)
        self.loaded_vectorstores = {}  # Cache des vector stores chargés

    def repair_all_vectorstores(self):
        """Tente de réparer tous les vector stores"""
        repaired = 0
        for path in self.base_dir.iterdir():
            if path.is_dir() and path.name.startswith("faiss_db_"):
                vectorstore_id = path.name.replace("faiss_db_", "")
                
                # Vérifier l'existence des fichiers essentiels
                if not (path / "index.faiss").exists():
                    continue
                    
                # Charger les métadonnées (les créera si manquantes)
                metadata = self.load_vectorstore_metadata(vectorstore_id)
                if metadata and metadata.get('is_auto_generated'):
                    repaired += 1
        
        if repaired > 0:
            logger.info(f"{repaired} vector stores réparés automatiquement")
        return repaired
    
    def generate_vectorstore_id(self, documents: List[LangchainDocument], 
                              embedding_model: str, chunk_params: Dict, custom_name: str = None) -> str:
        """Génère un ID unique pour le vector store basé sur le contenu et les paramètres"""
        content_hash = hashlib.md5()
        
        # Hash du contenu des documents
        for doc in documents:
            content_hash.update(doc.page_content.encode('utf-8'))
            # Inclure les métadonnées principales
            if 'source' in doc.metadata:
                content_hash.update(doc.metadata['source'].encode('utf-8'))
        
        # Hash des paramètres
        params_str = f"{embedding_model}_{chunk_params['chunk_size']}_{chunk_params['chunk_overlap']}"
        content_hash.update(params_str.encode('utf-8'))
        
        # Utiliser le nom personnalisé si fourni, sinon timestamp
        if custom_name:
            content_hash.update(custom_name.encode('utf-8'))
        
        timestamp = str(int(time.time()))
        content_hash.update(timestamp.encode('utf-8'))
        
        return content_hash.hexdigest()[:16]
    
    def get_vectorstore_path(self, vectorstore_id: str) -> Path:
        """Retourne le chemin du répertoire du vector store"""
        return self.base_dir / f"faiss_db_{vectorstore_id}"
    
    def vectorstore_exists(self, vectorstore_id: str) -> bool:
        """Vérifie si un vector store existe déjà avec vérification renforcée"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        
        # Vérifications multiples pour s'assurer de la validité
        if not vectorstore_path.exists():
            return False
            
        index_file = vectorstore_path / "index.faiss"
        pkl_file = vectorstore_path / "index.pkl"
        
        # Au minimum, l'index FAISS doit exister
        if not index_file.exists():
            logger.warning(f"Fichier index.faiss manquant pour {vectorstore_id}")
            return False
            
        # Vérifier la taille du fichier index
        if index_file.stat().st_size == 0:
            logger.warning(f"Fichier index.faiss vide pour {vectorstore_id}")
            return False
            
        return True
    
    def save_vectorstore_metadata(self, vectorstore_id: str, metadata: Dict):
        """Sauvegarde les métadonnées du vector store avec gestion d'erreurs"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            vectorstore_path.mkdir(parents=True, exist_ok=True)
            
            metadata_file = vectorstore_path / "metadata.json"
            
            # Sauvegarde avec backup
            temp_file = vectorstore_path / "metadata_temp.json"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False, default=str)
            
            # Remplacer le fichier principal seulement si la sauvegarde temp réussit
            if temp_file.exists():
                temp_file.replace(metadata_file)
                logger.info(f"Métadonnées sauvegardées pour {vectorstore_id}")
            
        except Exception as e:
            logger.error(f"Erreur sauvegarde métadonnées {vectorstore_id}: {str(e)}")
    
    def load_vectorstore_metadata(self, vectorstore_id: str) -> Optional[Dict]:
        """Charge les métadonnées avec gestion robuste des erreurs"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            metadata_file = vectorstore_path / "metadata.json"
            
            # Vérifier l'existence et la taille du fichier
            if not metadata_file.exists():
                logger.debug(f"Fichier metadata.json manquant pour {vectorstore_id}")
                return self._create_basic_metadata(vectorstore_id)
                
            if metadata_file.stat().st_size == 0:
                logger.warning(f"Fichier metadata.json vide pour {vectorstore_id}")
                return self._create_basic_metadata(vectorstore_id)
            
            # Charger avec validation
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
                
            # Validation des champs obligatoires
            required_fields = ['vectorstore_id', 'embedding_model', 'created_at']
            if not all(field in metadata for field in required_fields):
                logger.warning(f"Métadonnées incomplètes pour {vectorstore_id}")
                return self._create_basic_metadata(vectorstore_id)
                
            return metadata
            
        except json.JSONDecodeError:
            logger.error(f"Metadata corrompues (JSON invalide) pour {vectorstore_id}")
            return self._create_basic_metadata(vectorstore_id)
        except Exception as e:
            logger.error(f"Erreur critique chargement metadata {vectorstore_id}: {str(e)}")
            return None

    def _create_basic_metadata(self, vectorstore_id: str) -> Dict:
        """Crée des métadonnées de base pour un vector store"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        created_at = datetime.fromtimestamp(vectorstore_path.stat().st_ctime).isoformat()
        
        basic_metadata = {
            'vectorstore_id': vectorstore_id,
            'custom_name': f"VS_{vectorstore_id[:8]}",
            'created_at': created_at,
            'embedding_model': 'unknown',
            'document_count': 0,
            'vector_count': 0,
            'is_auto_generated': True
        }
        
        # Sauvegarder pour les prochains chargements
        self.save_vectorstore_metadata(vectorstore_id, basic_metadata)
        return basic_metadata
    
    def load_vectorstore(self, vectorstore_id: str, embeddings) -> Optional[FAISS]:
        """Charge un vector store FAISS depuis le disque"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            if not self.vectorstore_exists(vectorstore_id):
                logger.error(f"Vector store {vectorstore_id} n'existe pas ou est corrompu")
                return None
            
            # Charger les métadonnées d'abord pour vérifier la compatibilité
            metadata = self.load_vectorstore_metadata(vectorstore_id)
            if not metadata:
                logger.error(f"Impossible de charger les métadonnées pour {vectorstore_id}")
                return None
            
            # Charger le vector store
            try:
                vectorstore = FAISS.load_local(
                    str(vectorstore_path), 
                    embeddings,
                    allow_dangerous_deserialization=True
                )
        
                # Mettre en cache avec les métadonnées
                self.loaded_vectorstores[vectorstore_id] = {
                    'vectorstore': vectorstore,
                    'metadata': metadata,
                    'embedding_model': metadata.get('embedding_model', 'unknown')
                }
                
                logger.info(f"Vector store chargé avec succès: {vectorstore_id}")
                return vectorstore
                
            except Exception as e:
                logger.error(f"Erreur lors du chargement FAISS {vectorstore_id}: {str(e)}")
                # Nettoyer le cache en cas d'erreur
                if vectorstore_id in self.loaded_vectorstores:
                    del self.loaded_vectorstores[vectorstore_id]
                return None
                
        except Exception as e:
            logger.error(f"Erreur chargement vector store {vectorstore_id}: {str(e)}")
            return None
        
    def save_vectorstore(self, vectorstore: FAISS, vectorstore_id: str, 
                        documents: List[LangchainDocument], embedding_model: str,
                        chunk_params: Dict, custom_name: str = None) -> bool:
        """Sauvegarde le vector store avec gestion d'erreurs renforcée"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            # Supprimer le répertoire existant si nécessaire
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
            
            # Créer le répertoire
            vectorstore_path.mkdir(parents=True, exist_ok=True)
            
            # Sauvegarder le vector store avec gestion d'erreur
            try:
                vectorstore.save_local(str(vectorstore_path))
            except Exception as e:
                logger.error(f"Erreur sauvegarde FAISS {vectorstore_id}: {str(e)}")
                return False
            
            # Vérifier que les fichiers ont été créés
            index_file = vectorstore_path / "index.faiss"
            pkl_file = vectorstore_path / "index.pkl"
            
            if not index_file.exists() or index_file.stat().st_size == 0:
                logger.error(f"Fichier index.faiss invalide pour {vectorstore_id}")
                return False
            
            # Générer un nom personnalisé si non fourni
            if not custom_name:
                sources = list(set([doc.metadata.get('source', 'Unknown')[:20] for doc in documents[:3]]))
                custom_name = f"VS_{len(sources)}docs_{datetime.now().strftime('%Y%m%d_%H%M')}"
            
            # Créer les métadonnées complètes
            metadata = {
                'vectorstore_id': vectorstore_id,
                'custom_name': custom_name,
                'created_at': datetime.now().isoformat(),
                'embedding_model': embedding_model,
                'chunk_params': chunk_params,
                'document_count': len(documents),
                'vector_count': vectorstore.index.ntotal if hasattr(vectorstore.index, 'ntotal') else len(documents),
                'documents_info': [],
                'sources_mapping': {},
                'version': "1.0"  # Version pour compatibilité future
            }
            
            # Ajouter les informations détaillées des documents
            doc_sources = {}
            for doc in documents:
                source = doc.metadata.get('source', 'Unknown')
                doc_type = doc.metadata.get('type', 'unknown')
                doc_id = doc.metadata.get('doc_id', 'unknown')
                chunk_id = doc.metadata.get('chunk_id', 'unknown')
                
                if source not in doc_sources:
                    doc_sources[source] = {
                        'source': source,
                        'type': doc_type,
                        'chunks': 0,
                        'total_size': 0,
                        'doc_ids': set(),
                        'first_chunk_preview': doc.page_content[:200]
                    }
                
                doc_sources[source]['chunks'] += 1
                doc_sources[source]['total_size'] += len(doc.page_content)
                doc_sources[source]['doc_ids'].add(doc_id)
                
                # Mapping pour traçabilité
                metadata['sources_mapping'][chunk_id] = {
                    'source': source,
                    'doc_type': doc_type,
                    'doc_id': doc_id,
                    'metadata': doc.metadata
                }
            
            # Convertir les sets en listes pour JSON
            for source_info in doc_sources.values():
                source_info['doc_ids'] = list(source_info['doc_ids'])
            
            metadata['documents_info'] = list(doc_sources.values())
            
            # Sauvegarder les métadonnées
            self.save_vectorstore_metadata(vectorstore_id, metadata)
            
            # Vérifier que les métadonnées ont été sauvegardées
            if not (vectorstore_path / "metadata.json").exists():
                logger.error(f"Métadonnées non sauvegardées pour {vectorstore_id}")
                return False
            
            # Mettre en cache
            self.loaded_vectorstores[vectorstore_id] = {
                'vectorstore': vectorstore,
                'metadata': metadata,
                'embedding_model': embedding_model
            }
            
            logger.info(f"Vector store sauvegardé avec succès: {vectorstore_path}")
            return True
            
        except Exception as e:
            logger.error(f"Erreur sauvegarde vector store {vectorstore_id}: {str(e)}")
            # Nettoyer en cas d'erreur
            if vectorstore_path.exists():
                try:
                    shutil.rmtree(vectorstore_path)
                except:
                    pass
            return False
    
    def list_vectorstores(self) -> List[Dict]:
        """Liste tous les vector stores disponibles avec validation"""
        vectorstores = []
        
        if not self.base_dir.exists():
            return vectorstores
        
        for path in self.base_dir.iterdir():
            if path.is_dir() and path.name.startswith("faiss_db_"):
                vectorstore_id = path.name.replace("faiss_db_", "")
                
                # Vérifier d'abord si le vector store est valide
                if not self.vectorstore_exists(vectorstore_id):
                    logger.warning(f"Vector store corrompu ignoré: {vectorstore_id}")
                    continue
                
                metadata = self.load_vectorstore_metadata(vectorstore_id)
                
                if metadata:
                    try:
                        # Ajouter des informations sur la taille
                        total_size = sum(f.stat().st_size for f in path.rglob('*') if f.is_file())
                        metadata['storage_size_mb'] = round(total_size / (1024 * 1024), 2)
                        metadata['path'] = str(path)
                        metadata['is_loaded'] = vectorstore_id in self.loaded_vectorstores
                        metadata['last_modified'] = datetime.fromtimestamp(path.stat().st_mtime).isoformat()
                        metadata['id'] = vectorstore_id  # Ajouter l'ID
                        
                        # Validation des données critiques
                        if 'embedding_model' not in metadata:
                            metadata['embedding_model'] = 'unknown'
                        if 'document_count' not in metadata:
                            metadata['document_count'] = 0
                        if 'vector_count' not in metadata:
                            metadata['vector_count'] = 0
                        
                        # Extraire les sources depuis documents_info
                        sources = []
                        if 'documents_info' in metadata:
                            for doc_info in metadata['documents_info']:
                                sources.append(doc_info.get('source', 'Unknown'))
                        metadata['sources'] = sources
                        metadata['num_documents'] = metadata.get('document_count', 0)
                        metadata['num_chunks'] = metadata.get('vector_count', 0)
                            
                        vectorstores.append(metadata)
                    except Exception as e:
                        logger.error(f"Erreur traitement metadata pour {vectorstore_id}: {str(e)}")
                        continue
                else:
                    # Créer des métadonnées minimales pour les vector stores sans metadata
                    logger.info(f"Métadonnées manquantes pour {vectorstore_id}, création de métadonnées de base")
                    try:
                        total_size = sum(f.stat().st_size for f in path.rglob('*') if f.is_file())
                        basic_metadata = {
                            'vectorstore_id': vectorstore_id,
                            'id': vectorstore_id,
                            'custom_name': f"VS_Legacy_{vectorstore_id[:8]}",
                            'created_at': datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
                            'embedding_model': 'unknown',
                            'document_count': 0,
                            'vector_count': 0,
                            'storage_size_mb': round(total_size / (1024 * 1024), 2),
                            'path': str(path),
                            'is_loaded': vectorstore_id in self.loaded_vectorstores,
                            'is_legacy': True,
                            'sources': [],
                            'num_documents': 0,
                            'num_chunks': 0
                        }
                        vectorstores.append(basic_metadata)
                    except Exception as e:
                        logger.error(f"Erreur création métadonnées de base pour {vectorstore_id}: {str(e)}")
        
        # Trier par date de création (plus récent d'abord)
        vectorstores.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        return vectorstores
    
    def delete_vectorstore(self, vectorstore_id: str) -> bool:
        """Supprime un vector store avec vérifications"""
        try:
            # Supprimer du cache
            if vectorstore_id in self.loaded_vectorstores:
                del self.loaded_vectorstores[vectorstore_id]
            
            # Supprimer du disque
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
                logger.info(f"Vector store supprimé: {vectorstore_id}")
                return True
            else:
                logger.warning(f"Vector store {vectorstore_id} n'existe pas sur le disque")
                return False
        except Exception as e:
            logger.error(f"Erreur suppression vector store {vectorstore_id}: {str(e)}")
            return False
    
    def get_active_vectorstores(self) -> Dict[str, Dict]:
        """Retourne tous les vector stores actuellement chargés"""
        return self.loaded_vectorstores.copy()
    
    def unload_vectorstore(self, vectorstore_id: str):
        """Décharge un vector store du cache"""
        if vectorstore_id in self.loaded_vectorstores:
            del self.loaded_vectorstores[vectorstore_id]
            logger.info(f"Vector store déchargé du cache: {vectorstore_id}")
    
    def cleanup_vectorstores(self) -> int:
        """Nettoie tous les vector stores"""
        deleted_count = 0
        try:
            if self.base_dir.exists():
                for path in self.base_dir.iterdir():
                    if path.is_dir() and path.name.startswith("faiss_db_"):
                        shutil.rmtree(path)
                        deleted_count += 1
            
            # Vider le cache
            self.loaded_vectorstores.clear()
            
            logger.info(f"{deleted_count} vector stores supprimés")
        except Exception as e:
            logger.error(f"Erreur nettoyage: {str(e)}")
        
        return deleted_count


class DocumentProcessor:
    """Classe pour traiter différents types de documents avec métadonnées"""
    
    def __init__(self):
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        # Dictionnaire pour stocker tous les documents par leur ID unique
        self.document_registry = {}
    
    def reset(self):
        """Remet à zéro le processor pour traiter de nouveaux documents"""
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        self.document_registry = {}
    
    def create_metadata(self, source: str, doc_type: str, size: int = 0, 
                       additional_info: Dict = None) -> Dict:
        """Crée les métadonnées pour un document avec ID unique"""
        doc_id = f"{doc_type}_{hashlib.md5(f'{source}_{size}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"

        
        metadata = {
            'doc_id': doc_id,
            'source': source,
            'type': doc_type,
            'size': size,
            'processed_at': datetime.now().isoformat(),
            'hash': hashlib.md5(source.encode()).hexdigest()[:8]
        }
        if additional_info:
            metadata.update(additional_info)
        return metadata
    
    def register_document(self, doc_metadata: Dict, content: str):
        """Enregistre un document dans le registre global"""
        doc_id = doc_metadata['doc_id']
        self.document_registry[doc_id] = {
            'metadata': doc_metadata,
            'content': content,
            'chunks': []
        }
    
    def process_pdf(self, pdf_file) -> List[LangchainDocument]:
        """Traite un fichier PDF avec métadonnées"""
        documents = []
        try:
            pdf_reader = PdfReader(pdf_file)
            full_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    full_text += page_text + "\n\n"
                    
                    # Créer un document par page avec métadonnées
                    metadata = self.create_metadata(
                        source=pdf_file.name,
                        doc_type="pdf",
                        size=len(page_text),
                        additional_info={'page_number': page_num + 1}
                    )
                    
                    # Enregistrer dans le registre
                    self.register_document(metadata, page_text)
                    
                    documents.append(LangchainDocument(
                        page_content=page_text,
                        metadata=metadata
                    ))
            
            self.stats['file_types']['pdf'] = self.stats['file_types'].get('pdf', 0) + 1
            logger.info(f"PDF traité avec succès: {pdf_file.name} ({len(documents)} pages)")
            
        except Exception as e:
            error_msg = f"Erreur PDF {pdf_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_docx(self, docx_file) -> List[LangchainDocument]:
        """Traite un fichier DOCX avec métadonnées"""
        documents = []
        try:
            # Essayer d'abord avec docx2txt
            text = docx2txt.process(docx_file)
            if not text.strip():
                # Fallback avec python-docx
                doc = Document(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if text.strip():
                metadata = self.create_metadata(
                    source=docx_file.name,
                    doc_type="docx",
                    size=len(text),
                    additional_info={'word_count': len(text.split())}
                )
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['docx'] = self.stats['file_types'].get('docx', 0) + 1
                logger.info(f"DOCX traité avec succès: {docx_file.name}")
            
        except Exception as e:
            error_msg = f"Erreur DOCX {docx_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_txt(self, txt_file) -> List[LangchainDocument]:
        """Traite un fichier TXT avec métadonnées"""
        documents = []
        try:
            # Détecter l'encodage
            raw_data = txt_file.read()
            if isinstance(raw_data, bytes):
                try:
                    text = raw_data.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text = raw_data.decode('latin-1')
                    except UnicodeDecodeError:
                        text = raw_data.decode('cp1252', errors='ignore')
            else:
                text = raw_data
            
            if text.strip():
                metadata = self.create_metadata(
                    source=txt_file.name,
                    doc_type="txt",
                    size=len(text),
                    additional_info={'word_count': len(text.split())}
                )
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['txt'] = self.stats['file_types'].get('txt', 0) + 1
                logger.info(f"TXT traité avec succès: {txt_file.name}")
            
        except Exception as e:
            error_msg = f"Erreur TXT {txt_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_excel(self, excel_file) -> List[LangchainDocument]:
        """Traite un fichier Excel avec métadonnées améliorées"""
        documents = []
        try:
            # Détecter le type de fichier Excel
            file_extension = excel_file.name.lower().split('.')[-1]
            
            if file_extension == 'xlsx':
                df_dict = pd.read_excel(excel_file, engine='openpyxl', sheet_name=None)
            elif file_extension == 'xls':
                df_dict = pd.read_excel(excel_file, engine='xlrd', sheet_name=None)
            else:
                df_dict = pd.read_excel(excel_file, sheet_name=None)
            
            total_sheets = len(df_dict)
            logger.info(f"Traitement de {total_sheets} feuilles Excel dans {excel_file.name}")
            
            # Traiter chaque feuille séparément
            for sheet_name, df in df_dict.items():
                text_content = []
                
                # En-tête avec informations sur la feuille
                text_content.append(f"Feuille: {sheet_name}")
                text_content.append(f"Colonnes: {', '.join(df.columns.tolist())}")
                text_content.append(f"Nombre de lignes: {len(df)}")
                
                # Traitement des colonnes
                for column in df.columns:
                    column_data = df[column].dropna()
                    if not column_data.empty:
                        text_content.append(f"\nColonne {column}:")
                        
                        # Ajouter les statistiques pour les colonnes numériques
                        if pd.api.types.is_numeric_dtype(df[column]):
                            stats = df[column].describe()
                            text_content.append(f"Type: Numérique")
                            text_content.append(f"Moyenne: {stats['mean']:.2f}")
                            text_content.append(f"Minimum: {stats['min']:.2f}")
                            text_content.append(f"Maximum: {stats['max']:.2f}")
                            text_content.append(f"Médiane: {stats['50%']:.2f}")
                        else:
                            text_content.append(f"Type: Texte/Autre")
                            # Pour les colonnes non-numériques, afficher les valeurs uniques
                            unique_values = column_data.unique()
                            text_content.append(f"Valeurs uniques: {len(unique_values)}")
                        
                        # Limiter à 50 valeurs par colonne pour éviter des documents trop lourds
                        values = column_data.head(50).tolist()
                        values_str = ", ".join(str(v) for v in values)
                        text_content.append(f"Échantillon de valeurs: {values_str}")
                        if len(column_data) > 50:
                            text_content.append(f"... et {len(column_data) - 50} autres valeurs")

                # Créer un document pour cette feuille
                text = "\n".join(text_content)
                
                metadata = self.create_metadata(
                    source=f"{excel_file.name}::{sheet_name}",
                    doc_type="excel",
                    size=len(text),
                    additional_info={
                        'sheet_name': sheet_name,
                        'total_sheets': total_sheets,
                        'rows': len(df),
                        'columns': len(df.columns),
                        'column_names': df.columns.tolist(),
                        'has_numeric_data': any(pd.api.types.is_numeric_dtype(df[col]) for col in df.columns)
                    }
                )
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
            
            self.stats['file_types']['excel'] = self.stats['file_types'].get('excel', 0) + 1
            logger.info(f"Excel traité avec succès: {excel_file.name} ({total_sheets} feuilles)")
            
        except Exception as e:
            error_msg = f"Erreur Excel {excel_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_csv(self, csv_file) -> List[LangchainDocument]:
        """Traite un fichier CSV avec métadonnées"""
        documents = []
        try:
            # Essayer différents encodages et séparateurs
            encodings = ['utf-8', 'latin-1', 'cp1252']
            separators = [',', ';', '\t']
            
            df = None
            used_encoding = None
            used_separator = None
            
            for encoding in encodings:
                for separator in separators:
                    try:
                        df = pd.read_csv(csv_file, encoding=encoding, sep=separator)
                        if len(df.columns) > 1:  # Vérifier si la séparation a fonctionné
                            used_encoding = encoding
                            used_separator = separator
                            break
                    except:
                        continue
                if df is not None:
                    break
            
            if df is None:
                raise ValueError("Impossible de lire le fichier CSV avec les encodages/séparateurs testés")
            
            # Convertir DataFrame en texte
            text_content = []
            text_content.append(f"Fichier CSV: {csv_file.name}")
            text_content.append(f"Encodage: {used_encoding}, Séparateur: '{used_separator}'")
            text_content.append(f"Colonnes: {', '.join(df.columns.tolist())}")
            text_content.append(f"Nombre de lignes: {len(df)}")
            text_content.append("\nDonnées:")
            
            # Ajouter les données avec limite
            max_rows = 1000
            df_sample = df.head(max_rows)
            text_content.append(df_sample.to_string(index=False))
            
            if len(df) > max_rows:
                text_content.append(f"\n... ({len(df) - max_rows} lignes supplémentaires tronquées)")
            
            text = "\n".join(text_content)
            
            metadata = self.create_metadata(
                source=csv_file.name,
                doc_type="csv",
                size=len(text),
                additional_info={
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'encoding': used_encoding,
                    'separator': used_separator
                }
            )
            
            # Enregistrer dans le registre
            self.register_document(metadata, text)
            
            documents.append(LangchainDocument(
                page_content=text,
                metadata=metadata
            ))
            
            self.stats['file_types']['csv'] = self.stats['file_types'].get('csv', 0) + 1
            logger.info(f"CSV traité avec succès: {csv_file.name} ({len(df)} lignes)")
            
        except Exception as e:
            error_msg = f"Erreur CSV {csv_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_web_url(self, url: str) -> List[LangchainDocument]:
        """Traite une URL web avec métadonnées"""
        documents = []
        try:
            # Valider l'URL
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            # Charger le contenu web
            loader = WebBaseLoader(url)
            web_docs = loader.load()
            
            for doc in web_docs:
                if doc.page_content.strip():
                    metadata = self.create_metadata(
                        source=url,
                        doc_type="web",
                        size=len(doc.page_content),
                        additional_info={
                            'url': url,
                            'title': doc.metadata.get('title', 'Sans titre'),
                            'word_count': len(doc.page_content.split())
                        }
                    )
                    
                    # Enregistrer dans le registre
                    self.register_document(metadata, doc.page_content)
                    
                    documents.append(LangchainDocument(
                        page_content=doc.page_content,
                        metadata=metadata
                    ))
            
            self.stats['file_types']['web'] = self.stats['file_types'].get('web', 0) + 1
            logger.info(f"URL traitée avec succès: {url}")
            
        except Exception as e:
            error_msg = f"Erreur URL {url}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_files(self, files, urls=None) -> List[LangchainDocument]:
        """Traite une liste de fichiers et URLs"""
        start_time = time.time()
        all_documents = []
        
        # Traiter les fichiers
        if files:
            self.stats['total_docs'] += len(files)
            for file in files:
                try:
                    file_extension = file.name.lower().split('.')[-1]
                    
                    if file_extension == 'pdf':
                        docs = self.process_pdf(file)
                    elif file_extension == 'docx':
                        docs = self.process_docx(file)
                    elif file_extension == 'txt':
                        docs = self.process_txt(file)
                    elif file_extension in ['xlsx', 'xls']:
                        docs = self.process_excel(file)
                    elif file_extension == 'csv':
                        docs = self.process_csv(file)
                    else:
                        raise ValueError(f"Type de fichier non supporté: {file_extension}")
                    
                    all_documents.extend(docs)
                    self.processed_docs.extend(docs)
                    
                    if docs:
                        self.stats['successful_docs'] += 1
                    else:
                        self.stats['failed_docs'] += 1
                        
                except Exception as e:
                    error_msg = f"Erreur traitement {file.name}: {str(e)}"
                    self.errors.append(error_msg)
                    self.stats['failed_docs'] += 1
                    logger.error(error_msg)
        
        # Traiter les URLs
        if urls:
            url_list = [url.strip() for url in urls.split('\n') if url.strip()]
            self.stats['total_docs'] += len(url_list)
            
            for url in url_list:
                try:
                    docs = self.process_web_url(url)
                    all_documents.extend(docs)
                    self.processed_docs.extend(docs)
                    
                    if docs:
                        self.stats['successful_docs'] += 1
                    else:
                        self.stats['failed_docs'] += 1
                        
                except Exception as e:
                    error_msg = f"Erreur traitement URL {url}: {str(e)}"
                    self.errors.append(error_msg)
                    self.stats['failed_docs'] += 1
                    logger.error(error_msg)
        
        self.stats['processing_time'] = time.time() - start_time
        logger.info(f"Traitement terminé: {len(all_documents)} documents en {self.stats['processing_time']:.2f}s")
        
        return all_documents
    
    def get_processing_summary(self) -> Dict:
        """Retourne un résumé du traitement"""
        return {
            'stats': self.stats,
            'errors': self.errors,
            'total_chunks': len(self.processed_docs),
            'document_registry_size': len(self.document_registry)
        }


def get_text_chunks(documents: List[LangchainDocument], chunk_size: int = 1000, 
                   chunk_overlap: int = 200) -> List[LangchainDocument]:
    """Découpe les documents en chunks avec métadonnées préservées"""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunked_documents = []
        chunk_counter = 0
        
        for doc in documents:
            # Découper le document
            chunks = text_splitter.split_text(doc.page_content)
            
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    # Créer de nouvelles métadonnées pour le chunk
                    chunk_metadata = doc.metadata.copy()
                    chunk_metadata.update({
                        'chunk_id': f"{doc.metadata.get('doc_id', 'unknown')}_{i}",
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'chunk_size': len(chunk),
                        'global_chunk_id': chunk_counter
                    })
                    
                    chunked_documents.append(LangchainDocument(
                        page_content=chunk,
                        metadata=chunk_metadata
                    ))
                    chunk_counter += 1
        
        logger.info(f"Documents découpés: {len(documents)} -> {len(chunked_documents)} chunks")
        return chunked_documents
        
    except Exception as e:
        logger.error(f"Erreur découpage documents: {str(e)}")
        return documents


def get_embeddings(model_name: str):
    """Initialise le modèle d'embeddings"""
    try:
        if model_name not in EMBEDDING_MODELS:
            raise ValueError(f"Modèle d'embedding non supporté: {model_name}")
        
        model_config = EMBEDDING_MODELS[model_name]
        
        # Vérifier la clé API si nécessaire
        if model_config["requires_api_key"]:
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("Clé API OpenAI requise mais non trouvée")
            model_config["params"]["openai_api_key"] = api_key
        
        embeddings = model_config["class"](**model_config["params"])
        logger.info(f"Modèle d'embeddings initialisé: {model_name}")
        return embeddings
        
    except Exception as e:
        logger.error(f"Erreur initialisation embeddings {model_name}: {str(e)}")
        raise


def get_vectorstore(text_chunks: List[LangchainDocument], embeddings, 
                   vectorstore_manager: VectorStoreManager,
                   embedding_model: str, chunk_params: Dict,
                   custom_name: str = None, force_recreate: bool = False) -> Tuple[FAISS, str]:
    """Crée ou charge un vector store"""
    try:
        # Générer l'ID du vector store
        vectorstore_id = vectorstore_manager.generate_vectorstore_id(
            text_chunks, embedding_model, chunk_params, custom_name
        )
        
        # Vérifier si le vector store existe déjà
        if not force_recreate and vectorstore_manager.vectorstore_exists(vectorstore_id):
            logger.info(f"Chargement du vector store existant: {vectorstore_id}")
            vectorstore = vectorstore_manager.load_vectorstore(vectorstore_id, embeddings)
            if vectorstore:
                return vectorstore, vectorstore_id
            else:
                logger.warning(f"Échec du chargement, recréation du vector store: {vectorstore_id}")
        
        # Créer un nouveau vector store
        logger.info(f"Création d'un nouveau vector store: {vectorstore_id}")
        with st.spinner("Création du vector store en cours..."):
            vectorstore = FAISS.from_documents(text_chunks, embeddings)
        
        # Sauvegarder le vector store
        success = vectorstore_manager.save_vectorstore(
            vectorstore, vectorstore_id, text_chunks, 
            embedding_model, chunk_params, custom_name
        )
        
        if success:
            logger.info(f"Vector store créé et sauvegardé: {vectorstore_id}")
        else:
            logger.warning(f"Vector store créé mais échec de sauvegarde: {vectorstore_id}")
        
        return vectorstore, vectorstore_id
        
    except Exception as e:
        logger.error(f"Erreur création vector store: {str(e)}")
        raise


def get_conversation_chain(vectorstore, model_name: str):
    """Crée la chaîne de conversation"""
    try:
        if model_name not in LLM_MODELS:
            raise ValueError(f"Modèle LLM non supporté: {model_name}")
        
        model_config = LLM_MODELS[model_name]
        api_key = os.getenv(model_config["api_key_env"])
        
        if not api_key:
            raise ValueError(f"Clé API {model_config['api_key_env']} manquante")
        
        llm = ChatGoogleGenerativeAI(
            model=model_config["model_name"],
            google_api_key=api_key,
            temperature=0.3,
            convert_system_message_to_human=True
        )
        
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            ),
            memory=memory,
            return_source_documents=True,
            verbose=True
        )
        
        logger.info(f"Chaîne de conversation créée avec {model_name}")
        return conversation_chain
        
    except Exception as e:
        logger.error(f"Erreur création chaîne conversation: {str(e)}")
        raise


def handle_userinput(user_question, conversation_chain):
    """Traite l'input utilisateur et affiche la réponse"""
    try:
        with st.spinner("Génération de la réponse..."):
            response = conversation_chain({"question": user_question})
        
        # Stocker la conversation dans l'état de session
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []
        
        st.session_state.chat_history.append({
            "question": user_question,
            "answer": response["answer"],
            "source_documents": response.get("source_documents", []),
            "timestamp": datetime.now().isoformat()
        })
        
        # Afficher la conversation avec le template HTML
        for i, message in enumerate(st.session_state.chat_history):
            st.write(user_template.replace("{{MSG}}", message["question"]), 
                    unsafe_allow_html=True)
            st.write(bot_template.replace("{{MSG}}", message["answer"]), 
                    unsafe_allow_html=True)
            
            # Afficher les sources si disponibles
            if message["source_documents"]:
                with st.expander(f"📄 Sources utilisées ({len(message['source_documents'])})"):
                    for j, doc in enumerate(message["source_documents"]):
                        source = doc.metadata.get('source', 'Source inconnue')
                        doc_type = doc.metadata.get('type', 'unknown')
                        chunk_id = doc.metadata.get('chunk_id', 'unknown')
                        
                        st.write(f"**Source {j+1}:** {source} ({doc_type})")
                        st.write(f"**Chunk ID:** {chunk_id}")
                        
                        # Afficher un extrait du contenu
                        content_preview = doc.page_content[:300]
                        if len(doc.page_content) > 300:
                            content_preview += "..."
                        st.write(f"**Extrait:** {content_preview}")
                        st.write("---")
        
    except Exception as e:
        st.error(f"Erreur lors du traitement de la question: {str(e)}")
        logger.error(f"Erreur handle_userinput: {str(e)}")


def display_vectorstore_info(vectorstore_manager: VectorStoreManager):
    """Affiche les informations des vector stores"""
    st.subheader("📊 Gestion des Vector Stores")
    
    vectorstores = vectorstore_manager.list_vectorstores()
    
    if not vectorstores:
        st.info("Aucun vector store trouvé.")
        return
    
    # Statistiques générales
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Vector Stores", len(vectorstores))
    with col2:
        total_size = sum(vs.get('storage_size_mb', 0) for vs in vectorstores)
        st.metric("Taille totale", f"{total_size:.1f} MB")
    with col3:
        loaded_count = sum(1 for vs in vectorstores if vs.get('is_loaded', False))
        st.metric("Chargés", loaded_count)
    with col4:
        total_docs = sum(vs.get('document_count', 0) for vs in vectorstores)
        st.metric("Total documents", total_docs)
    
    # Liste des vector stores
    for vs in vectorstores:
        with st.expander(f"🗃️ {vs.get('custom_name', vs['vectorstore_id'])} - {vs.get('storage_size_mb', 0):.1f} MB"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.write(f"**ID:** {vs['vectorstore_id']}")
                st.write(f"**Créé le:** {vs.get('created_at', 'Inconnu')}")
                st.write(f"**Modèle d'embedding:** {vs.get('embedding_model', 'Inconnu')}")
                st.write(f"**Documents:** {vs.get('document_count', 0)}")
                st.write(f"**Chunks:** {vs.get('vector_count', 0)}")
                
                if vs.get('is_loaded'):
                    st.success("✅ Chargé en mémoire")
                else:
                    st.info("💾 Sur disque")
            
            with col2:
                # Actions
                if st.button(f"🗑️ Supprimer", key=f"delete_{vs['vectorstore_id']}"):
                    if vectorstore_manager.delete_vectorstore(vs['vectorstore_id']):
                        st.success("Vector store supprimé avec succès!")
                        st.rerun()
                    else:
                        st.error("Erreur lors de la suppression")
                
                if not vs.get('is_loaded'):
                    if st.button(f"📂 Charger", key=f"load_{vs['vectorstore_id']}"):
                        st.info("Chargement en cours...")
                        # Récupérer le modèle d'embedding approprié
                        embedding_model = vs.get('embedding_model', 'HuggingFace Sentence Transformers')
                        embeddings = get_embeddings(embedding_model)
                        
                        # Tenter le chargement
                        vectorstore = vectorstore_manager.load_vectorstore(vs['vectorstore_id'], embeddings)
                        if vectorstore:
                            # Créer la chaîne de conversation
                            st.session_state.conversation = get_conversation_chain(
                                vectorstore, 
                                list(LLM_MODELS.keys())[0]  # Utiliser le premier modèle LLM par défaut
                            )
                            st.session_state.current_vectorstore_id = vs['vectorstore_id']
                            st.success("Vector store chargé avec succès!")
                            
                            # Rediriger vers l'onglet chat
                            st.session_state.active_tab = "Chat"
                            st.rerun()
                        else:
                            st.error("Erreur lors du chargement du vector store")
                else:
                    if st.button(f"📤 Décharger", key=f"unload_{vs['vectorstore_id']}"):
                        vectorstore_manager.unload_vectorstore(vs['vectorstore_id'])
                        st.success("Vector store déchargé!")
                        st.rerun()
            
            # Afficher les sources si disponibles
            if vs.get('sources'):
                st.write("**Sources:**")
                for source in vs['sources'][:5]:  # Limiter à 5 sources
                    st.write(f"- {source}")
                if len(vs['sources']) > 5:
                    st.write(f"... et {len(vs['sources']) - 5} autres")
    
    # Actions globales
    st.write("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("🔧 Réparer les Vector Stores"):
            repaired = vectorstore_manager.repair_all_vectorstores()
            st.success(f"{repaired} vector stores réparés")
            st.rerun()
    
    with col2:
        if st.button("🗑️ Nettoyer tout", type="secondary"):
            if st.checkbox("Confirmer la suppression de tous les vector stores"):
                deleted = vectorstore_manager.cleanup_vectorstores()
                st.success(f"{deleted} vector stores supprimés")
                st.rerun()
    
    with col3:
        if st.button("🔄 Actualiser"):
            st.rerun()


def main():
    """Fonction principale de l'application Streamlit"""
    load_dotenv()
    
    st.set_page_config(
        page_title="Assistant de Documents IA", 
        page_icon="🤖",
        layout="wide"
    )
    
    st.write(css, unsafe_allow_html=True)
    st.header("🤖 Assistant de Documents avec IA")
    
    # Initialiser les objets dans session_state
    if "vectorstore_manager" not in st.session_state:
        st.session_state.vectorstore_manager = VectorStoreManager()
    
    if "document_processor" not in st.session_state:
        st.session_state.document_processor = DocumentProcessor()
    
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    
    if "current_vectorstore_id" not in st.session_state:
        st.session_state.current_vectorstore_id = None
    
    # Ajouter ici
    if "active_tab" not in st.session_state:
        st.session_state.active_tab = "Vector Stores"
    
    # Sidebar pour la configuration
    with st.sidebar:
        st.subheader("⚙️ Configuration")
        
        # Sélection du modèle d'embedding
        embedding_model = st.selectbox(
            "Modèle d'Embedding:",
            list(EMBEDDING_MODELS.keys()),
            index=2,  # HuggingFace Sentence Transformers par défaut
            help="Choisissez le modèle pour convertir le texte en vecteurs"
        )
        
        # Sélection du modèle LLM
        llm_model = st.selectbox(
            "Modèle de Langage:",
            list(LLM_MODELS.keys()),
            help="Choisissez le modèle pour générer les réponses"
        )
        
        # Paramètres de chunking
        st.subheader("🔧 Paramètres de découpage")
        chunk_size = st.slider("Taille des chunks", 500, 2000, 1000, 100)
        chunk_overlap = st.slider("Chevauchement", 0, 500, 200, 50)
        
        chunk_params = {
            'chunk_size': chunk_size,
            'chunk_overlap': chunk_overlap
        }
        
        # Upload de fichiers
        st.subheader("📁 Charger des documents")
        uploaded_files = st.file_uploader(
            "Choisissez vos fichiers",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt', 'csv', 'xlsx', 'xls'],
            help="Formats supportés: PDF, DOCX, TXT, CSV, XLSX, XLS"
        )
        
        # URLs
        st.subheader("🌐 URLs Web")
        urls_input = st.text_area(
            "URLs (une par ligne):",
            placeholder="https://example.com\nhttps://another-site.com",
            help="Entrez les URLs à traiter, une par ligne"
        )
        
        # Nom personnalisé pour le vector store
        custom_name = st.text_input(
            "Nom du Vector Store (optionnel):",
            placeholder="Mon projet de recherche",
            help="Donnez un nom personnalisé à votre vector store"
        )
        
        # Options de traitement
        force_recreate = st.checkbox(
            "Forcer la recréation",
            help="Force la création d'un nouveau vector store même s'il existe déjà"
        )
        
        # Bouton de traitement
        if st.button("🚀 Traiter les documents", type="primary"):
            if uploaded_files or urls_input.strip():
                try:
                    with st.spinner("Traitement en cours..."):
                        # Reset du processor
                        st.session_state.document_processor.reset()
                        
                        # Traitement des documents
                        documents = st.session_state.document_processor.process_files(
                            uploaded_files, urls_input.strip() if urls_input.strip() else None
                        )
                        
                        if documents:
                            # Découpage en chunks
                            text_chunks = get_text_chunks(documents, chunk_size, chunk_overlap)
                            
                            # Création des embeddings
                            embeddings = get_embeddings(embedding_model)
                            
                            # Création/chargement du vector store
                            vectorstore, vectorstore_id = get_vectorstore(
                                text_chunks, embeddings,
                                st.session_state.vectorstore_manager,
                                embedding_model, chunk_params,
                                custom_name, force_recreate
                            )
                            
                            # Création de la chaîne de conversation
                            st.session_state.conversation = get_conversation_chain(
                                vectorstore, llm_model
                            )
                            st.session_state.current_vectorstore_id = vectorstore_id
                            
                            # Afficher le résumé
                            summary = st.session_state.document_processor.get_processing_summary()
                            st.success(f"✅ Traitement terminé!")
                            st.write(f"📊 **Résumé:** {summary['stats']['successful_docs']} documents traités, "
                                   f"{summary['total_chunks']} chunks créés en {summary['stats']['processing_time']:.2f}s")
                            
                            if summary['errors']:
                                st.warning(f"⚠️ {len(summary['errors'])} erreurs détectées:")
                                for error in summary['errors']:
                                    st.write(f"- {error}")
                        else:
                            st.error("Aucun document n'a pu être traité.")
                            
                except Exception as e:
                    st.error(f"Erreur lors du traitement: {str(e)}")
                    logger.error(f"Erreur main traitement: {str(e)}")
            else:
                st.warning("Veuillez charger des fichiers ou entrer des URLs.")
    
    # Interface principale
    tab1, tab2, tab3 = st.tabs(["💬 Chat", "📊 Vector Stores", "ℹ️ Informations"])
    
    # Sélectionner automatiquement l'onglet chat si nécessaire
    if st.session_state.active_tab == "Chat":
        tab1.active = True
        st.session_state.active_tab = None  # Réinitialiser pour les prochains clics
    
    with tab1:
        # Interface de chat
        if st.session_state.conversation is not None:
            st.subheader("💬 Posez vos questions")
            
            user_question = st.text_input(
                "Votre question:",
                placeholder="Posez une question sur vos documents...",
                key="user_question"
            )
            
            if user_question:
                handle_userinput(user_question, st.session_state.conversation)
            
            # Afficher l'historique existant
            if "chat_history" in st.session_state and st.session_state.chat_history:
                st.write("---")
                st.subheader("📝 Historique de la conversation")
                for message in st.session_state.chat_history:
                    st.write(user_template.replace("{{MSG}}", message["question"]), 
                            unsafe_allow_html=True)
                    st.write(bot_template.replace("{{MSG}}", message["answer"]), 
                            unsafe_allow_html=True)
        else:
            st.info("👈 Veuillez d'abord traiter des documents dans la barre latérale pour commencer à poser des questions.")
            
            # Boutons d'exemples de questions
            if st.session_state.current_vectorstore_id:
                st.subheader("💡 Questions d'exemple")
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button("📋 Résume-moi les documents"):
                        st.session_state.user_question = "Peux-tu me faire un résumé des documents que tu as analysés ?"
                        st.rerun()
                    
                    if st.button("🔍 Quels sont les points clés ?"):
                        st.session_state.user_question = "Quels sont les points clés et les informations importantes dans ces documents ?"
                        st.rerun()
                
                with col2:
                    if st.button("📊 Y a-t-il des données chiffrées ?"):
                        st.session_state.user_question = "Y a-t-il des données chiffrées, des statistiques ou des métriques importantes ?"
                        st.rerun()
                    
                    if st.button("🎯 Quelles sont les conclusions ?"):
                        st.session_state.user_question = "Quelles sont les principales conclusions ou recommandations ?"
                        st.rerun()
    
    with tab2:
        # Gestion des Vector Stores
        display_vectorstore_info(st.session_state.vectorstore_manager)
    
    with tab3:
        # Informations sur l'application
        st.subheader("ℹ️ À propos de l'application")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("""
            **🎯 Fonctionnalités principales:**
            - 📁 Support multi-formats (PDF, DOCX, TXT, CSV, XLSX, XLS)
            - 🌐 Traitement d'URLs web
            - 🤖 Chat intelligent avec vos documents
            - 💾 Gestion persistante des vector stores
            - 🔧 Paramètres de chunking configurables
            - 📊 Statistiques détaillées
            """)
            
            st.write("""
            **🔧 Modèles disponibles:**
            - **Embeddings:** OpenAI, Google, HuggingFace
            - **LLM:** Google Gemini Pro/Flash
            - **Vector Store:** FAISS (Facebook AI)
            """)
        
        with col2:
            st.write("""
            **📝 Comment utiliser:**
            1. Configurez les modèles dans la barre latérale
            2. Chargez vos documents ou entrez des URLs
            3. Ajustez les paramètres de découpage si nécessaire
            4. Cliquez sur "Traiter les documents"
            5. Posez vos questions dans l'onglet Chat
            """)
            
            st.write("""
            **⚙️ Configuration requise:**
            - Clé API Google Gemini (GOOGLE_API_KEY)
            - Clé API OpenAI (optionnelle, pour embeddings OpenAI)
            - Python 3.8+ avec les dépendances installées
            """)
        
        # Statistiques de session
        if hasattr(st.session_state, 'document_processor'):
            summary = st.session_state.document_processor.get_processing_summary()
            
            st.subheader("📈 Statistiques de la session")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Documents traités", summary['stats']['successful_docs'])
            with col2:
                st.metric("Erreurs", summary['stats']['failed_docs'])
            with col3:
                st.metric("Chunks créés", summary['total_chunks'])
            with col4:
                processing_time = summary['stats'].get('processing_time', 0)
                st.metric("Temps de traitement", f"{processing_time:.2f}s")
            
            # Répartition par type de fichier
            if summary['stats']['file_types']:
                st.subheader("📊 Répartition par type de fichier")
                file_types_df = pd.DataFrame([
                    {'Type': k, 'Nombre': v} 
                    for k, v in summary['stats']['file_types'].items()
                ])
                st.bar_chart(file_types_df.set_index('Type'))
        
        # Logs et debugging
        with st.expander("🔍 Logs et Debugging"):
            if hasattr(st.session_state, 'document_processor') and st.session_state.document_processor.errors:
                st.write("**Erreurs rencontrées:**")
                for error in st.session_state.document_processor.errors:
                    st.code(error)
            else:
                st.info("Aucune erreur enregistrée.")
            
            # Informations sur l'environnement
            st.write("**Variables d'environnement:**")
            env_status = {
                'GOOGLE_API_KEY': '✅ Configurée' if os.getenv('GOOGLE_API_KEY') else '❌ Manquante',
                'OPENAI_API_KEY': '✅ Configurée' if os.getenv('OPENAI_API_KEY') else '⚠️ Optionnelle'
            }
            for key, status in env_status.items():
                st.write(f"- {key}: {status}")
    # Footer
    st.write("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; font-size: 0.8em;'>
            🤖 Assistant de Documents IA - Propulsé par Streamlit, LangChain et Google Gemini
        </div>
        """, 
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()
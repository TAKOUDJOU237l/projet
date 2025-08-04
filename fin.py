import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings, HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.document_loaders import WebBaseLoader
import requests
from bs4 import BeautifulSoup
import urllib.parse
import pandas as pd
import openpyxl
import xlrd
from pydantic import BaseModel
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument
from htmlTemplates import css, bot_template, user_template

import os
import warnings
import hashlib
import json
import shutil
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import logging
from pathlib import Path
import time

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Supprimer les avertissements
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Répertoires de stockage
VECTORSTORE_BASE_DIR = "vectorstores"
CACHE_DIR = "cache"

# Configuration des modèles
EMBEDDING_MODELS = {
    "OpenAI": {
        "class": OpenAIEmbeddings,
        "params": {},
        "requires_api_key": True,
        "display_name": "OpenAI Text-Embedding-3"
    },
    "HuggingFace Instructor": {
        "class": HuggingFaceInstructEmbeddings,
        "params": {"model_name": "hku-nlp/instructor-xl"},
        "requires_api_key": False,
        "display_name": "HF Instructor XL"
    },
    "HuggingFace Sentence Transformers": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/all-MiniLM-L6-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False,
        "display_name": "HF MiniLM-L6-v2"
    },
    "HuggingFace BGE Small": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "BAAI/bge-small-en-v1.5",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False,
        "display_name": "HF BGE Small v1.5"
    },
    "HuggingFace Multilingual": {
        "class": HuggingFaceEmbeddings,
        "params": {
            "model_name": "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            "model_kwargs": {"device": "cpu"}
        },
        "requires_api_key": False,
        "display_name": "HF Multilingual MiniLM"
    }
}

LLM_MODELS = {
    "Gemini 2.0 Flash": {
        "model_name": "gemini-2.0-flash-exp",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle puissant et rapide",
        "display_name": "Gemini 2.0 Flash"
    },
    "Gemini Pro 1.5": {
        "model_name": "gemini-1.5-pro-latest",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini 1.5 Pro - Modèle le plus avancé",
        "display_name": "Gemini 1.5 Pro"
    },
    "Gemini Pro": {
        "model_name": "gemini-pro",
        "api_key_env": "GOOGLE_API_KEY",
        "description": "Google Gemini Pro - Modèle standard performant",
        "display_name": "Gemini Pro"
    }
}

class VectorStoreManager:
    """Gestionnaire des vector stores FAISS avec sauvegarde en répertoire"""
    
    def __init__(self, base_dir: str = VECTORSTORE_BASE_DIR):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(exist_ok=True)
        self.cache_dir = Path(CACHE_DIR)
        self.cache_dir.mkdir(exist_ok=True)
        
    def generate_vectorstore_id(self, documents: List[LangchainDocument], 
                              embedding_model: str, chunk_params: Dict) -> str:
        """Génère un ID unique pour le vector store basé sur le contenu et les paramètres"""
        content_hash = hashlib.md5()
        
        # Hash du contenu des documents
        for doc in documents:
            content_hash.update(doc.page_content.encode('utf-8'))
            # Inclure les métadonnées principales
            if 'source' in doc.metadata:
                content_hash.update(doc.metadata['source'].encode('utf-8'))
        
        # Hash des paramètres
        params_str = f"{embedding_model}_{chunk_params['chunk_size']}_{chunk_params['chunk_overlap']}"
        content_hash.update(params_str.encode('utf-8'))
        
        # Ajouter timestamp pour unicité
        timestamp = str(int(time.time()))
        content_hash.update(timestamp.encode('utf-8'))
        
        return content_hash.hexdigest()[:16]
    
    def get_vectorstore_path(self, vectorstore_id: str) -> Path:
        """Retourne le chemin du répertoire du vector store"""
        return self.base_dir / f"faiss_db_{vectorstore_id}"
    
    def vectorstore_exists(self, vectorstore_id: str) -> bool:
        """Vérifie si un vector store existe déjà"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        return vectorstore_path.exists() and (vectorstore_path / "index.faiss").exists()
    
    def save_vectorstore_metadata(self, vectorstore_id: str, metadata: Dict):
        """Sauvegarde les métadonnées du vector store"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        metadata_file = vectorstore_path / "metadata.json"
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, indent=2, ensure_ascii=False, fp=f)
    
    def load_vectorstore_metadata(self, vectorstore_id: str) -> Optional[Dict]:
        """Charge les métadonnées du vector store"""
        vectorstore_path = self.get_vectorstore_path(vectorstore_id)
        metadata_file = vectorstore_path / "metadata.json"
        
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return None
    
    def save_vectorstore(self, vectorstore: FAISS, vectorstore_id: str, 
                        documents: List[LangchainDocument], embedding_model: str,
                        chunk_params: Dict, llm_model: str = None, custom_name: str = None) -> bool:
        """Sauvegarde le vector store en répertoire avec métadonnées détaillées"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            # Supprimer le répertoire existant si nécessaire
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
            
            # Créer le répertoire
            vectorstore_path.mkdir(parents=True, exist_ok=True)
            
            # Sauvegarder le vector store
            vectorstore.save_local(str(vectorstore_path))
            
            # Analyser les documents pour extraire les informations détaillées
            files_info = {}
            urls_info = []
            total_chunks = 0
            
            for doc in documents:
                source = doc.metadata.get('source', 'Unknown')
                doc_type = doc.metadata.get('type', 'unknown')
                
                if doc_type == 'web':
                    # Pour les URLs
                    url_info = {
                        'url': source,
                        'title': doc.metadata.get('title', 'Sans titre'),
                        'chunks': doc.metadata.get('total_chunks', 1),
                        'content_size': doc.metadata.get('size', 0),
                        'status_code': doc.metadata.get('status_code', 'unknown'),
                        'content_type': doc.metadata.get('content_type', 'unknown')
                    }
                    
                    # Vérifier si cette URL n'est pas déjà ajoutée
                    existing_url = next((u for u in urls_info if u['url'] == source), None)
                    if not existing_url:
                        urls_info.append(url_info)
                    
                else:
                    # Pour les fichiers
                    if source not in files_info:
                        files_info[source] = {
                            'filename': source,
                            'type': doc_type,
                            'pages': set() if doc_type == 'pdf' else None,
                            'sheets': set() if doc_type == 'excel' else None,
                            'total_chunks': 0,
                            'total_size': 0,
                            'word_count': 0
                        }
                    
                    # Mettre à jour les informations
                    files_info[source]['total_chunks'] += 1
                    files_info[source]['total_size'] += doc.metadata.get('size', 0)
                    
                    if doc_type == 'pdf' and 'page_number' in doc.metadata:
                        files_info[source]['pages'].add(doc.metadata['page_number'])
                    elif doc_type == 'excel' and 'sheet_name' in doc.metadata:
                        files_info[source]['sheets'].add(doc.metadata['sheet_name'])
                    elif doc_type == 'docx' and 'word_count' in doc.metadata:
                        files_info[source]['word_count'] = doc.metadata['word_count']
                
                total_chunks += 1
            
            # Convertir les sets en listes pour la sérialisation JSON
            for file_info in files_info.values():
                if file_info['pages'] is not None:
                    file_info['pages'] = sorted(list(file_info['pages']))
                if file_info['sheets'] is not None:
                    file_info['sheets'] = sorted(list(file_info['sheets']))
            
            # Générer un nom personnalisé si non fourni
            if not custom_name:
                file_count = len(files_info)
                url_count = len(urls_info)
                if file_count > 0 and url_count > 0:
                    custom_name = f"Mix_{file_count}fichiers_{url_count}urls_{datetime.now().strftime('%m%d_%H%M')}"
                elif file_count > 0:
                    if file_count == 1:
                        filename = list(files_info.keys())[0]
                        # Extraire le nom sans extension
                        name_without_ext = Path(filename).stem
                        custom_name = f"{name_without_ext}_{datetime.now().strftime('%m%d_%H%M')}"
                    else:
                        custom_name = f"{file_count}fichiers_{datetime.now().strftime('%m%d_%H%M')}"
                elif url_count > 0:
                    custom_name = f"{url_count}urls_{datetime.now().strftime('%m%d_%H%M')}"
                else:
                    custom_name = f"VS_{datetime.now().strftime('%Y%m%d_%H%M')}"
            
            # Créer les métadonnées complètes
            metadata = {
                'vectorstore_id': vectorstore_id,
                'custom_name': custom_name,
                'created_at': datetime.now().isoformat(),
                'embedding_model': embedding_model,
                'embedding_model_display': EMBEDDING_MODELS.get(embedding_model, {}).get('display_name', embedding_model),
                'llm_model': llm_model,
                'llm_model_display': LLM_MODELS.get(llm_model, {}).get('display_name', llm_model) if llm_model else None,
                'chunk_params': chunk_params,
                'document_count': len(set([doc.metadata.get('source') for doc in documents])),
                'total_chunks': total_chunks,
                'vector_count': vectorstore.index.ntotal,
                'files_info': list(files_info.values()),
                'urls_info': urls_info,
                'processing_stats': {
                    'total_files': len(files_info),
                    'total_urls': len(urls_info),
                    'total_size_mb': round(sum([f['total_size'] for f in files_info.values()]) / (1024 * 1024), 2),
                    'file_types': list(set([f['type'] for f in files_info.values()]))
                }
            }
            
            # Sauvegarder les métadonnées
            self.save_vectorstore_metadata(vectorstore_id, metadata)
            
            logger.info(f"Vector store sauvegardé: {vectorstore_path}")
            return True
            
        except Exception as e:
            logger.error(f"Erreur sauvegarde vector store {vectorstore_id}: {str(e)}")
            return False
    
    def load_vectorstore(self, vectorstore_id: str, embeddings) -> Optional[FAISS]:
        """Charge un vector store depuis le répertoire"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            
            if not self.vectorstore_exists(vectorstore_id):
                logger.warning(f"Vector store {vectorstore_id} n'existe pas")
                return None
            
            # Charger le vector store
            vectorstore = FAISS.load_local(
                str(vectorstore_path), 
                embeddings, 
                allow_dangerous_deserialization=True
            )
            
            logger.info(f"Vector store chargé: {vectorstore_path}")
            return vectorstore
            
        except Exception as e:
            logger.error(f"Erreur chargement vector store {vectorstore_id}: {str(e)}")
            return None
    
    def list_vectorstores(self) -> List[Dict]:
        """Liste tous les vector stores disponibles avec leurs métadonnées"""
        vectorstores = []
        
        for path in self.base_dir.iterdir():
            if path.is_dir() and path.name.startswith("faiss_db_"):
                vectorstore_id = path.name.replace("faiss_db_", "")
                metadata = self.load_vectorstore_metadata(vectorstore_id)
                
                if metadata:
                    # Ajouter des informations sur la taille de stockage
                    total_size = sum(path.stat().st_size for path in path.rglob('*') if path.is_file())
                    metadata['storage_size_mb'] = round(total_size / (1024 * 1024), 2)
                    metadata['path'] = str(path)
                    vectorstores.append(metadata)
        
        # Trier par date de création (plus récent d'abord)
        vectorstores.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        return vectorstores
    
    def delete_vectorstore(self, vectorstore_id: str) -> bool:
        """Supprime un vector store"""
        try:
            vectorstore_path = self.get_vectorstore_path(vectorstore_id)
            if vectorstore_path.exists():
                shutil.rmtree(vectorstore_path)
                logger.info(f"Vector store supprimé: {vectorstore_id}")
                return True
            return False
        except Exception as e:
            logger.error(f"Erreur suppression vector store {vectorstore_id}: {str(e)}")
            return False
    
    def cleanup_old_vectorstores(self, max_age_days: int = 30, max_count: int = 10):
        """Nettoie les anciens vector stores"""
        vectorstores = self.list_vectorstores()
        current_time = datetime.now()
        deleted_count = 0
        
        # Supprimer les vector stores trop anciens
        for vs in vectorstores:
            created_at = datetime.fromisoformat(vs['created_at'])
            age_days = (current_time - created_at).days
            
            if age_days > max_age_days:
                self.delete_vectorstore(vs['vectorstore_id'])
                deleted_count += 1
        
        # Garder seulement les N plus récents
        if len(vectorstores) > max_count:
            for vs in vectorstores[max_count:]:
                self.delete_vectorstore(vs['vectorstore_id'])
                deleted_count += 1
        
        if deleted_count > 0:
            logger.info(f"Nettoyage: {deleted_count} vector stores supprimés")

class DocumentProcessor:
    """Classe pour traiter différents types de documents avec métadonnées"""
    
    def __init__(self):
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        # Dictionnaire pour stocker tous les documents par leur ID unique
        self.document_registry = {}
    
    def reset(self):
        """Remet à zéro le processor pour traiter de nouveaux documents"""
        self.processed_docs = []
        self.errors = []
        self.stats = {
            'total_docs': 0,
            'successful_docs': 0,
            'failed_docs': 0,
            'processing_time': 0,
            'file_types': {}
        }
        self.document_registry = {}
    
    def create_metadata(self, source: str, doc_type: str, size: int = 0, 
                       additional_info: Dict = None) -> Dict:
        """Crée les métadonnées pour un document avec ID unique"""
        doc_id = f"{doc_type}_{hashlib.md5(f'{source}_{size}_{datetime.now().isoformat()}'.encode()).hexdigest()[:12]}"

        
        metadata = {
            'doc_id': doc_id,
            'source': source,
            'type': doc_type,
            'size': size,
            'processed_at': datetime.now().isoformat(),
            'hash': hashlib.md5(source.encode()).hexdigest()[:8]
        }
        if additional_info:
            metadata.update(additional_info)
        return metadata
    
    def register_document(self, doc_metadata: Dict, content: str):
        """Enregistre un document dans le registre global"""
        doc_id = doc_metadata['doc_id']
        self.document_registry[doc_id] = {
            'metadata': doc_metadata,
            'content': content,
            'chunks': []
        }
    
    def process_pdf(self, pdf_file) -> List[LangchainDocument]:
        """Traite un fichier PDF avec métadonnées"""
        documents = []
        try:
            pdf_reader = PdfReader(pdf_file)
            full_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    full_text += page_text + "\n\n"
                    
                    # Créer un document par page avec métadonnées
                    metadata = self.create_metadata(
                        source=pdf_file.name,
                        doc_type="pdf",
                        size=len(page_text),
                        additional_info={'page_number': page_num + 1}
                    )
                    
                    # Enregistrer dans le registre
                    self.register_document(metadata, page_text)
                    
                    documents.append(LangchainDocument(
                        page_content=page_text,
                        metadata=metadata
                    ))
            
            self.stats['file_types']['pdf'] = self.stats['file_types'].get('pdf', 0) + 1
            logger.info(f"PDF traité avec succès: {pdf_file.name} ({len(documents)} pages)")
            
        except Exception as e:
            error_msg = f"Erreur PDF {pdf_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_docx(self, docx_file) -> List[LangchainDocument]:
        """Traite un fichier DOCX avec métadonnées"""
        documents = []
        try:
            # Essayer d'abord avec docx2txt
            text = docx2txt.process(docx_file)
            if not text.strip():
                # Fallback avec python-docx
                doc = Document(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if text.strip():
                metadata = self.create_metadata(
                    source=docx_file.name,
                    doc_type="docx",
                    size=len(text),
                    additional_info={'word_count': len(text.split())}
                )
                
                # Enregistrer dans le registre
                self.register_document(metadata, text)
                
                documents.append(LangchainDocument(
                    page_content=text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['docx'] = self.stats['file_types'].get('docx', 0) + 1
                logger.info(f"DOCX traité avec succès: {docx_file.name}")
            
        except Exception as e:
            error_msg = f"Erreur DOCX {docx_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_excel(self, excel_file) -> List[LangchainDocument]:
        """Traite un fichier Excel avec métadonnées"""
        documents = []
        try:
            # Déterminer le moteur selon l'extension
            engine = 'openpyxl' if excel_file.name.endswith('.xlsx') else 'xlrd'
            df_dict = pd.read_excel(excel_file, sheet_name=None, engine=engine)
            
            for sheet_name, df in df_dict.items():
                if not df.empty:
                    # Créer un texte structuré pour chaque feuille
                    text_content = f"=== Feuille: {sheet_name} ===\n\n"
                    
                    # Ajouter les en-têtes
                    headers = list(df.columns)
                    text_content += "Colonnes: " + " | ".join(str(h) for h in headers) + "\n\n"
                    
                    # Ajouter les données
                    for index, row in df.iterrows():
                        row_data = []
                        for col in headers:
                            cell_value = row[col]
                            if pd.notna(cell_value):
                                row_data.append(f"{col}: {cell_value}")
                        
                        if row_data:
                            text_content += f"Ligne {index + 1} - " + " | ".join(row_data) + "\n"
                    
                    # Créer le document avec métadonnées
                    metadata = self.create_metadata(
                        source=excel_file.name,
                        doc_type="excel",
                        size=len(text_content),
                        additional_info={
                            'sheet_name': sheet_name,
                            'rows': len(df),
                            'columns': len(df.columns)
                        }
                    )
                    
                    # Enregistrer dans le registre
                    self.register_document(metadata, text_content)
                    
                    documents.append(LangchainDocument(
                        page_content=text_content,
                        metadata=metadata
                    ))
            
            self.stats['file_types']['excel'] = self.stats['file_types'].get('excel', 0) + 1
            logger.info(f"Excel traité avec succès: {excel_file.name} ({len(documents)} feuilles)")
            
        except Exception as e:
            error_msg = f"Erreur Excel {excel_file.name}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def process_url(self, url: str) -> List[LangchainDocument]:
        """Traite une URL avec métadonnées"""
        documents = []
        try:
            if not self.is_valid_url(url):
                raise ValueError("URL invalide")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            start_time = time.time()
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Parser le contenu
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Supprimer les éléments indésirables
            for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
                tag.decompose()
            
            # Extraire le titre
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "Sans titre"
            
            # Extraire le texte principal
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if clean_text.strip():
                processing_time = time.time() - start_time
                
                metadata = self.create_metadata(
                    source=url,
                    doc_type="web",
                    size=len(clean_text),
                    additional_info={
                        'title': title_text,
                        'status_code': response.status_code,
                        'processing_time': round(processing_time, 2),
                        'content_type': response.headers.get('content-type', 'unknown')
                    }
                )
                
                # Enregistrer dans le registre
                self.register_document(metadata, clean_text)
                
                documents.append(LangchainDocument(
                    page_content=clean_text,
                    metadata=metadata
                ))
                
                self.stats['file_types']['web'] = self.stats['file_types'].get('web', 0) + 1
                logger.info(f"URL traitée avec succès: {url} ({len(clean_text)} caractères)")
                
        except Exception as e:
            error_msg = f"Erreur URL {url}: {str(e)}"
            self.errors.append(error_msg)
            logger.error(error_msg)
            
        return documents
    
    def is_valid_url(self, url: str) -> bool:
        """Valide une URL"""
        try:
            parsed = urllib.parse.urlparse(url)
            return all([parsed.scheme, parsed.netloc])
        except:
            return False
    
    def process_all_sources(self, uploaded_files: List = None, urls: List[str] = None) -> List[LangchainDocument]:
        """Traite toutes les sources avec gestion d'erreurs et statistiques"""
        start_time = time.time()
        all_documents = []
        
        self.stats['total_docs'] = len(uploaded_files or []) + len(urls or [])
        
        # Traiter les fichiers uploadés
        if uploaded_files:
            for file in uploaded_files:
                try:
                    if file.type == "application/pdf":
                        docs = self.process_pdf(file)
                    elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                      "application/msword"]:
                        docs = self.process_docx(file)
                    elif file.type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                      "application/vnd.ms-excel"]:
                        docs = self.process_excel(file)
                    else:
                        continue
                    
                    all_documents.extend(docs)
                    if docs:
                        self.stats['successful_docs'] += 1
                    else:
                        self.stats['failed_docs'] += 1
                        
                except Exception as e:
                    self.stats['failed_docs'] += 1
                    logger.error(f"Erreur traitement fichier {file.name}: {str(e)}")
        
        # Traiter les URLs
        if urls:
            for url in urls:
                try:
                    docs = self.process_url(url)
                    all_documents.extend(docs)
                    if docs:
                        self.stats['successful_docs'] += 1
                    else:
                        self.stats['failed_docs'] += 1
                        
                except Exception as e:
                    self.stats['failed_docs'] += 1
                    logger.error(f"Erreur traitement URL {url}: {str(e)}")
        
        self.stats['processing_time'] = round(time.time() - start_time, 2)
        self.processed_docs = all_documents
        
        logger.info(f"Traitement terminé: {self.stats['successful_docs']}/{self.stats['total_docs']} sources traitées")
        return all_documents

def get_text_chunks(documents: List[LangchainDocument], chunk_size: int = 1000, 
                   chunk_overlap: int = 200) -> List[LangchainDocument]:
    """Divise les documents en chunks avec préservation des métadonnées"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    all_chunks = []
    for doc in documents:
        try:
            # Diviser le document en chunks
            doc_chunks = text_splitter.split_text(doc.page_content)
            
            # Créer des documents LangChain avec métadonnées enrichies
            for i, chunk_text in enumerate(doc_chunks):
                # Copier les métadonnées originales
                chunk_metadata = doc.metadata.copy()
                
                # Ajouter des informations sur le chunk
                chunk_metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(doc_chunks),
                    'chunk_size': len(chunk_text),
                    'chunk_id': f"{doc.metadata.get('doc_id', 'unknown')}_{i}"
                })
                
                all_chunks.append(LangchainDocument(
                    page_content=chunk_text,
                    metadata=chunk_metadata
                ))
                
        except Exception as e:
            logger.error(f"Erreur division en chunks pour {doc.metadata.get('source', 'unknown')}: {str(e)}")
    
    logger.info(f"Documents divisés en {len(all_chunks)} chunks")
    return all_chunks

def get_embeddings(model_name: str):
    """Crée une instance d'embeddings selon le modèle choisi"""
    try:
        model_config = EMBEDDING_MODELS.get(model_name)
        if not model_config:
            raise ValueError(f"Modèle d'embedding non supporté: {model_name}")
        
        # Vérifier la clé API si nécessaire
        if model_config.get("requires_api_key", False):
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("Clé API OpenAI non trouvée dans les variables d'environnement")
        
        # Créer l'instance d'embeddings
        embeddings = model_config["class"](**model_config["params"])
        
        logger.info(f"Modèle d'embedding initialisé: {model_config['display_name']}")
        return embeddings
        
    except Exception as e:
        logger.error(f"Erreur initialisation embeddings {model_name}: {str(e)}")
        raise

def get_vectorstore(text_chunks: List[LangchainDocument], embeddings, 
                   vs_manager: VectorStoreManager, embedding_model: str,
                   chunk_params: Dict, llm_model: str = None, 
                   custom_name: str = None) -> Tuple[FAISS, str]:
    """Crée ou charge un vector store FAISS avec gestion de cache"""
    try:
        # Générer l'ID du vector store
        vectorstore_id = vs_manager.generate_vectorstore_id(text_chunks, embedding_model, chunk_params)
        
        # Essayer de charger depuis le cache
        if vs_manager.vectorstore_exists(vectorstore_id):
            logger.info(f"Chargement du vector store depuis le cache: {vectorstore_id}")
            vectorstore = vs_manager.load_vectorstore(vectorstore_id, embeddings)
            if vectorstore:
                return vectorstore, vectorstore_id
        
        # Créer un nouveau vector store
        logger.info(f"Création d'un nouveau vector store: {vectorstore_id}")
        with st.spinner("Création du vector store..."):
            vectorstore = FAISS.from_documents(text_chunks, embeddings)
        
        # Sauvegarder
        vs_manager.save_vectorstore(
            vectorstore, vectorstore_id, text_chunks, 
            embedding_model, chunk_params, llm_model, custom_name
        )
        
        return vectorstore, vectorstore_id
        
    except Exception as e:
        logger.error(f"Erreur création vector store: {str(e)}")
        raise

def get_conversation_chain(vectorstore: FAISS, llm_model: str):
    """Crée une chaîne de conversation avec le vector store"""
    try:
        model_config = LLM_MODELS.get(llm_model)
        if not model_config:
            raise ValueError(f"Modèle LLM non supporté: {llm_model}")
        
        # Vérifier la clé API
        api_key = os.getenv(model_config["api_key_env"])
        if not api_key:
            raise ValueError(f"Clé API {model_config['api_key_env']} non trouvée")
        
        # Créer le modèle LLM
        llm = ChatGoogleGenerativeAI(
            model=model_config["model_name"],
            google_api_key=api_key,
            temperature=0.1,
            convert_system_message_to_human=True
        )
        
        # Créer la mémoire de conversation
        memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        # Créer la chaîne de conversation
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            ),
            memory=memory,
            return_source_documents=True,
            verbose=True
        )
        
        logger.info(f"Chaîne de conversation créée avec {model_config['display_name']}")
        return conversation_chain
        
    except Exception as e:
        logger.error(f"Erreur création chaîne de conversation: {str(e)}")
        raise

def handle_userinput(user_question: str):
    """Gère la question de l'utilisateur et affiche la réponse"""
    if st.session_state.conversation is None:
        st.error("Veuillez d'abord traiter vos documents.")
        return
    
    try:
        with st.spinner("Génération de la réponse..."):
            response = st.session_state.conversation({'question': user_question})
        
        # Mettre à jour l'historique des messages
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        
        st.session_state.chat_history.append({
            'question': user_question,
            'answer': response['answer'],
            'source_documents': response.get('source_documents', []),
            'timestamp': datetime.now().isoformat()
        })
        
        # Afficher l'historique
        display_chat_history()
        
    except Exception as e:
        st.error(f"Erreur lors du traitement de la question: {str(e)}")
        logger.error(f"Erreur handle_userinput: {str(e)}")

def display_chat_history():
    """Affiche l'historique des conversations"""
    if 'chat_history' not in st.session_state or not st.session_state.chat_history:
        return
    
    # Utiliser un timestamp pour rendre les clés uniques
    current_time = datetime.now().timestamp()
    
    for i, chat in enumerate(reversed(st.session_state.chat_history)):
        # Question de l'utilisateur
        st.write(user_template.replace("{{MSG}}", chat['question']), unsafe_allow_html=True)
        
        # Réponse du bot
        st.write(bot_template.replace("{{MSG}}", chat['answer']), unsafe_allow_html=True)
        
        # Sources (dans un expander)
        if chat.get('source_documents'):
            with st.expander(f"📚 Sources pour la question {len(st.session_state.chat_history) - i}", expanded=False):
                for j, doc in enumerate(chat['source_documents']):
                    source = doc.metadata.get('source', 'Source inconnue')
                    doc_type = doc.metadata.get('type', 'unknown')
                    
                    # Informations contextuelles selon le type
                    context_info = ""
                    if doc_type == 'pdf' and 'page_number' in doc.metadata:
                        context_info = f" (Page {doc.metadata['page_number']})"
                    elif doc_type == 'excel' and 'sheet_name' in doc.metadata:
                        context_info = f" (Feuille: {doc.metadata['sheet_name']})"
                    elif doc_type == 'web' and 'title' in doc.metadata:
                        context_info = f" ({doc.metadata['title']})"
                    
                    st.write(f"**Source {j+1}:** {source}{context_info}")
                    
                    # Générer une clé unique pour chaque text_area
                    unique_key = f"source_{current_time}_{i}_{j}"
                    
                    # Afficher un extrait du contenu
                    content_preview = doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content
                    st.text_area(
                        f"Extrait {j+1}:", 
                        content_preview, 
                        height=100, 
                        key=unique_key,
                        disabled=True
                    )

def display_vectorstore_info(vs_manager: VectorStoreManager, vectorstore_id: str = None):
    """Affiche les informations détaillées sur le vector store"""
    if not vectorstore_id:
        return
    
    metadata = vs_manager.load_vectorstore_metadata(vectorstore_id)
    if not metadata:
        return
    
    with st.expander("📊 Informations sur le Vector Store", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            st.write(f"**Nom:** {metadata.get('custom_name', 'N/A')}")
            st.write(f"**ID:** {vectorstore_id}")
            st.write(f"**Créé le:** {metadata.get('created_at', 'N/A')[:16]}")
            st.write(f"**Modèle d'embedding:** {metadata.get('embedding_model_display', 'N/A')}")
            st.write(f"**Modèle LLM:** {metadata.get('llm_model_display', 'N/A')}")
        
        with col2:
            st.write(f"**Nombre de documents:** {metadata.get('document_count', 0)}")
            st.write(f"**Chunks totaux:** {metadata.get('total_chunks', 0)}")
            st.write(f"**Vecteurs:** {metadata.get('vector_count', 0)}")
            st.write(f"**Taille chunk:** {metadata.get('chunk_params', {}).get('chunk_size', 'N/A')}")
            st.write(f"**Overlap:** {metadata.get('chunk_params', {}).get('chunk_overlap', 'N/A')}")
        
        # Statistiques de traitement
        processing_stats = metadata.get('processing_stats', {})
        if processing_stats:
            st.write("**Statistiques de traitement:**")
            st.write(f"- Fichiers: {processing_stats.get('total_files', 0)}")
            st.write(f"- URLs: {processing_stats.get('total_urls', 0)}")
            st.write(f"- Taille totale: {processing_stats.get('total_size_mb', 0)} MB")
            st.write(f"- Types: {', '.join(processing_stats.get('file_types', []))}")
        
        # Détails des fichiers
        files_info = metadata.get('files_info', [])
        if files_info:
            st.write("**Fichiers traités:**")
            for file_info in files_info:
                filename = file_info.get('filename', 'N/A')
                file_type = file_info.get('type', 'unknown')
                chunks = file_info.get('total_chunks', 0)
                
                additional_info = ""
                if file_type == 'pdf' and 'pages' in file_info:
                    pages = file_info['pages']
                    if pages:
                        additional_info = f" ({len(pages)} pages)"
                elif file_type == 'excel' and 'sheets' in file_info:
                    sheets = file_info['sheets']
                    if sheets:
                        additional_info = f" ({len(sheets)} feuilles)"
                elif file_type == 'docx' and 'word_count' in file_info:
                    word_count = file_info['word_count']
                    additional_info = f" ({word_count} mots)"
                
                st.write(f"- {filename} ({file_type}){additional_info} - {chunks} chunks")
        
        # Détails des URLs
        urls_info = metadata.get('urls_info', [])
        if urls_info:
            st.write("**URLs traitées:**")
            for url_info in urls_info:
                url = url_info.get('url', 'N/A')
                title = url_info.get('title', 'Sans titre')
                chunks = url_info.get('chunks', 0)
                status = url_info.get('status_code', 'unknown')
                
                st.write(f"- {title} ({status}) - {chunks} chunks")
                st.write(f"  URL: {url}")

def main():
    """Fonction principale de l'application Streamlit"""
    # Configuration de la page
    st.set_page_config(
        page_title="Chat RAG Avancé",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Charger les variables d'environnement
    load_dotenv()
    
    # CSS personnalisé
    st.write(css, unsafe_allow_html=True)
    
    # Titre principal
    st.title("🤖 Chat RAG Avancé avec Gestion de Vector Stores")
    st.markdown("*Analysez vos documents et URLs avec une IA conversationnelle avancée*")
    
    # Initialiser les composants
    if 'vs_manager' not in st.session_state:
        st.session_state.vs_manager = VectorStoreManager()
    
    if 'doc_processor' not in st.session_state:
        st.session_state.doc_processor = DocumentProcessor()
    
    # Sidebar pour la configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Configuration des modèles
        st.subheader("🔧 Modèles")
        
        selected_embedding = st.selectbox(
            "Modèle d'embedding",
            options=list(EMBEDDING_MODELS.keys()),
            format_func=lambda x: EMBEDDING_MODELS[x]['display_name'],
            index=2  # Par défaut HuggingFace Sentence Transformers
        )
        
        selected_llm = st.selectbox(
            "Modèle de langage",
            options=list(LLM_MODELS.keys()),
            format_func=lambda x: LLM_MODELS[x]['display_name'],
            index=0  # Par défaut Gemini 2.0 Flash
        )
        
        # Configuration du chunking
        st.subheader("📝 Paramètres de découpage")
        chunk_size = st.slider("Taille des chunks", 500, 2000, 1000, 100)
        chunk_overlap = st.slider("Chevauchement", 50, 500, 200, 50)
        
        chunk_params = {
            'chunk_size': chunk_size,
            'chunk_overlap': chunk_overlap
        }
        
        st.divider()
        
        # Gestion des Vector Stores existants
        st.subheader("💾 Vector Stores Sauvegardés")
        
        vectorstores = st.session_state.vs_manager.list_vectorstores()
        
        if vectorstores:
            # Sélection d'un vector store existant
            vs_options = {f"{vs['custom_name']} ({vs['created_at'][:10]})": vs['vectorstore_id'] 
                         for vs in vectorstores}
            
            selected_vs_name = st.selectbox(
                "Charger un Vector Store existant",
                options=[""] + list(vs_options.keys()),
                format_func=lambda x: "Sélectionner..." if x == "" else x
            )
            
            if selected_vs_name and selected_vs_name != "":
                selected_vs_id = vs_options[selected_vs_name]
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🔄 Charger", key="load_vs"):
                        try:
                            embeddings = get_embeddings(selected_embedding)
                            vectorstore = st.session_state.vs_manager.load_vectorstore(selected_vs_id, embeddings)
                            
                            if vectorstore:
                                st.session_state.conversation = get_conversation_chain(vectorstore, selected_llm)
                                st.session_state.current_vs_id = selected_vs_id
                                st.success("Vector Store chargé avec succès!")
                                st.rerun()
                            else:
                                st.error("Erreur lors du chargement")
                                
                        except Exception as e:
                            st.error(f"Erreur: {str(e)}")
                
                with col2:
                    if st.button("🗑️ Supprimer", key="delete_vs"):
                        if st.session_state.vs_manager.delete_vectorstore(selected_vs_id):
                            st.success("Vector Store supprimé!")
                            st.rerun()
                        else:
                            st.error("Erreur lors de la suppression")
            
            # Afficher la liste des vector stores
            with st.expander("📋 Détails des Vector Stores", expanded=False):
                for vs in vectorstores:
                    st.write(f"**{vs['custom_name']}**")
                    st.write(f"- Créé: {vs['created_at'][:16]}")
                    st.write(f"- Documents: {vs['document_count']}")
                    st.write(f"- Chunks: {vs.get('total_chunks', 'Non disponible')}")

                    st.write(f"- Taille: {vs.get('storage_size_mb', 0)} MB")
                    st.write("---")
        else:
            st.info("Aucun Vector Store sauvegardé")
        
        # Nettoyage automatique
        if st.button("🧹 Nettoyer les anciens VS"):
            st.session_state.vs_manager.cleanup_old_vectorstores(max_age_days=7, max_count=5)
            st.success("Nettoyage effectué!")
            st.rerun()
    
    # Zone principale
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Interface de chat
        st.subheader("💬 Conversation")
        
        # Afficher l'état actuel
        if 'conversation' in st.session_state and st.session_state.conversation:
            if 'current_vs_id' in st.session_state:
                display_vectorstore_info(st.session_state.vs_manager, st.session_state.current_vs_id)
            
            st.success("✅ Prêt pour les questions!")
            
            # Zone de saisie de question
            user_question = st.text_input("Posez votre question:", key="user_question")
            
            if user_question:
                handle_userinput(user_question)
        else:
            st.info("Veuillez charger un Vector Store existant ou traiter de nouveaux documents.")
        
        # Affichage de l'historique des conversations
        if 'chat_history' in st.session_state and st.session_state.chat_history:
            st.divider()
            display_chat_history()
    
    with col2:
        # Interface de traitement des documents
        st.subheader("📄 Traitement des Documents")
        
        # Upload de fichiers
        uploaded_files = st.file_uploader(
            "Choisissez vos fichiers",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'doc', 'xlsx', 'xls']
        )
        
        # Saisie d'URLs
        st.subheader("🌐 URLs à traiter")
        urls_text = st.text_area(
            "URLs (une par ligne)",
            height=100,
            placeholder="https://example.com\nhttps://autre-site.com"
        )
        
        urls = [url.strip() for url in urls_text.split('\n') if url.strip()]
        
        # Nom personnalisé pour le vector store
        custom_vs_name = st.text_input(
            "Nom du Vector Store (optionnel)",
            placeholder="MonAnalyse_2024"
        )
        
        # Bouton de traitement
        if st.button("🚀 Traiter et Créer Vector Store", type="primary"):
            if not uploaded_files and not urls:
                st.error("Veuillez fournir au moins un fichier ou une URL.")
            else:
                try:
                    with st.spinner("Traitement en cours..."):
                        # Reset du processor
                        st.session_state.doc_processor.reset()
                        
                        # Traitement des documents
                        documents = st.session_state.doc_processor.process_all_sources(
                            uploaded_files=uploaded_files,
                            urls=urls
                        )
                        
                        if not documents:
                            st.error("Aucun document n'a pu être traité.")
                        else:
                            # Affichage des statistiques
                            stats = st.session_state.doc_processor.stats
                            st.success(f"✅ {stats['successful_docs']}/{stats['total_docs']} sources traitées")
                            
                            if st.session_state.doc_processor.errors:
                                with st.expander("⚠️ Erreurs de traitement", expanded=False):
                                    for error in st.session_state.doc_processor.errors:
                                        st.error(error)
                            
                            # Initialiser les embeddings avant de les utiliser
                            embeddings = get_embeddings(selected_embedding)
                            
                            # Division en chunks
                            text_chunks = get_text_chunks(documents, chunk_size, chunk_overlap)
                            vectorstore, vs_id = get_vectorstore(
                                text_chunks, embeddings, st.session_state.vs_manager,
                                selected_embedding, chunk_params, selected_llm, custom_vs_name
                            )
                            
                            # Création de la chaîne de conversation
                            st.session_state.conversation = get_conversation_chain(vectorstore, selected_llm)
                            st.session_state.current_vs_id = vs_id
                            
                            st.success("🎉 Vector Store créé et prêt pour les questions!")
                            st.rerun()
        
                except Exception as e:
                    st.error(f"Erreur lors du traitement: {str(e)}")
                    logger.error(f"Erreur main processing: {str(e)}")
        
        # Informations sur les fichiers uploadés
        if uploaded_files:
            with st.expander("📁 Fichiers sélectionnés", expanded=True):
                for file in uploaded_files:
                    file_size = len(file.getvalue()) / (1024 * 1024)  # MB
                    st.write(f"• {file.name} ({file.type}) - {file_size:.2f} MB")
        
        # Informations sur les URLs
        if urls:
            with st.expander("🔗 URLs sélectionnées", expanded=True):
                for url in urls:
                    st.write(f"• {url}")

if __name__ == "__main__":
    main()